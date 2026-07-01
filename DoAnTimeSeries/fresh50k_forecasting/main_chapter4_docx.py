from __future__ import annotations

import csv
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TABLE_DIR = BASE_DIR / "outputs" / "tables"
FIGURE_DIR = BASE_DIR / "outputs" / "figures"
OUT_DIR = BASE_DIR / "deliverables" / "reports"
OUT_PATH = OUT_DIR / "Chuong_4_Ket_qua_khoi_phuc_nhu_cau_va_du_bao.docx"
IMAGE_OUT_DIR = OUT_DIR / "Chuong_4_Hinh_anh_goc"


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
        doc.save(fallback)
        return fallback


def set_default_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12 if style_name == "Normal" else 14)
        style.font.bold = style_name.startswith("Heading")
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)

    for style_name, italic, bold in [("Caption", True, False), ("Table Caption", False, True)]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.font.italic = italic
        style.font.bold = bold
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    doc.add_paragraph(caption, style="Table Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, export_filename: str, width: float = 6.2) -> None:
    path = FIGURE_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[Thiếu hình: {filename}]")
        return
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, IMAGE_OUT_DIR / export_filename)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")


def read_metric_table(filename: str) -> dict[str, str]:
    result: dict[str, str] = {}
    with (TABLE_DIR / filename).open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            result[row["Metric"]] = row["Value"]
    return result


def fmt_float(value: str | float, digits: int = 4) -> str:
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)


def fmt_pct(value: str | float, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def format_cell(col: str, value) -> str:
    if isinstance(value, float):
        pct_cols = ["WAPE", "sMAPE", "WPE", "lift", "share"]
        if any(key.lower() in col.lower() for key in pct_cols):
            return fmt_pct(value)
        return fmt_float(value, 4)
    return str(value)


def df_rows(filename: str, columns: list[str], max_rows: int | None = None) -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / filename)
    if max_rows is not None:
        df = df.head(max_rows)
    rows = []
    for _, row in df.iterrows():
        rows.append([format_cell(col, row[col]) for col in columns])
    return rows


def selected_forecast_rows() -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_forecasting_comparison.csv")
    order = [
        "Observed-sales forecasting",
        "Recovered seasonal naive 7-day",
        "Recovered-demand forecasting",
        "Recovered seasonal-ML hybrid",
    ]
    df = df[df["Evaluation target"] == "Recovered latent demand proxy"].copy()
    df["rank_key"] = df["Model"].map({name: idx for idx, name in enumerate(order)})
    df = df[df["rank_key"].notna()].sort_values("rank_key")
    cols = ["Model", "RMSE", "MAE", "WAPE", "sMAPE", "WPE"]
    return [[format_cell(col, row[col]) for col in cols] for _, row in df.iterrows()]


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if IMAGE_OUT_DIR.exists():
        shutil.rmtree(IMAGE_OUT_DIR)
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)

    recovery = read_metric_table("owner_latent_recovery_summary.csv")
    uplift = read_metric_table("imputation_uplift_summary.csv")
    forecast = pd.read_csv(TABLE_DIR / "owner_two_stage_forecasting_comparison.csv")
    pseudo = pd.read_csv(TABLE_DIR / "imputation_pseudo_stockout_aggregate_validation.csv")
    cap_sensitivity = pd.read_csv(TABLE_DIR / "imputation_cap_sensitivity.csv")
    rec_target = forecast[forecast["Evaluation target"] == "Recovered latent demand proxy"].copy()
    observed_wape = float(rec_target[rec_target["Model"] == "Observed-sales forecasting"]["WAPE"].iloc[0])
    hybrid_wape = float(rec_target[rec_target["Model"] == "Recovered seasonal-ML hybrid"]["WAPE"].iloc[0])
    seasonal_wape = float(rec_target[rec_target["Model"] == "Recovered seasonal naive 7-day"]["WAPE"].iloc[0])
    recovered_wape = float(rec_target[rec_target["Model"] == "Recovered-demand forecasting"]["WAPE"].iloc[0])
    rel_improve = (observed_wape - hybrid_wape) / observed_wape
    pseudo_hourly_wape = float(pseudo[pseudo["Validation level"] == "Hourly aggregate"]["WAPE"].iloc[0])
    pseudo_daily_wape = float(pseudo[pseudo["Validation level"] == "Daily aggregate"]["WAPE"].iloc[0])
    pseudo_ratio = float(pseudo[pseudo["Validation level"] == "Daily aggregate"]["Prediction / actual ratio"].iloc[0])
    cap_lookup = cap_sensitivity.set_index("Scenario")
    q50_lift = float(cap_lookup.loc["Cap lost-demand at q50", "Recovered lift over observed"])
    q90_lift = float(cap_lookup.loc["Cap lost-demand at q90", "Recovered lift over observed"])
    q90_share = float(cap_lookup.loc["Cap lost-demand at q90", "Share of original recovered lost demand"])
    q100_lift = float(cap_lookup.loc["Cap lost-demand at q100", "Recovered lift over observed"])
    sub_velocity_path = TABLE_DIR / "stockout_substitution_velocity_diagnostics.csv"
    sub_velocity_rows = (
        df_rows(
            "stockout_substitution_velocity_diagnostics.csv",
            [
                "Segment",
                "Rows",
                "Mean sale velocity ratio 3h/24h",
                "Mean sale momentum 3h-6h",
                "Mean peer sales same group",
                "Mean peer velocity ratio 3h/24h",
                "Mean peer stockout rate",
            ],
        )
        if sub_velocity_path.exists()
        else []
    )

    doc = Document()
    set_default_font(doc)

    title = doc.add_heading("CHƯƠNG 4: KẾT QUẢ KHÔI PHỤC NHU CẦU VÀ DỰ BÁO", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        "Chương 4 trình bày kết quả thực nghiệm của pipeline hai giai đoạn đã mô tả ở Chương 3. "
        "Phần đầu đánh giá liệu cơ chế recovery có khôi phục một lượng nhu cầu hợp lý tại các giờ stockout hay không. "
        "Phần sau đánh giá liệu forecast trên recovered demand có cải thiện so với forecast trực tiếp observed sales, đồng thời phân tích vai trò của seasonal baseline và hybrid seasonal-ML.",
    )
    add_paragraph(
        doc,
        "Mục tiêu của chương không chỉ là báo cáo chỉ số WAPE/RMSE, mà còn kiểm tra tính hợp lý của toàn bộ quy trình: imputation có làm tổng demand tăng quá mức không, seasonal naive có phải baseline mạnh không, và mô hình học máy đóng góp thêm gì ngoài pattern tuần.",
    )

    doc.add_heading("4.1. Đánh giá Phục hồi Nhu cầu Tiềm ẩn", level=2)
    doc.add_heading("4.1.1. Tổng quan Kết quả Khôi phục Nhu cầu", level=3)
    add_paragraph(
        doc,
        "Kết quả recovery cho thấy observed sales thấp hơn recovered latent demand, phù hợp với giả thuyết rằng stockout làm doanh số ghi nhận bị kiểm duyệt. "
        f"Tổng observed sales đạt {fmt_float(recovery.get('Observed sales', ''), 2)}, trong khi recovered latent demand đạt {fmt_float(recovery.get('Recovered latent demand', ''), 2)}. "
        f"Phần recovered lost demand là {fmt_float(recovery.get('Recovered lost demand', ''), 2)}, tương ứng mức tăng {fmt_pct(recovery.get('Recovered lift over observed', ''))} so với observed sales.",
    )
    add_table(
        doc,
        ["Chỉ tiêu", "Giá trị"],
        [
            ["Observed sales", fmt_float(recovery.get("Observed sales", ""), 2)],
            ["Recovered latent demand", fmt_float(recovery.get("Recovered latent demand", ""), 2)],
            ["Recovered lost demand", fmt_float(recovery.get("Recovered lost demand", ""), 2)],
            ["Recovered lift over observed", fmt_pct(recovery.get("Recovered lift over observed", ""))],
            ["Stockout row rate", fmt_pct(recovery.get("Stockout row rate", ""))],
            ["Recovery method", recovery.get("Recovery method", "")],
        ],
        "Bảng 4.1: Tổng quan kết quả latent demand recovery",
    )
    add_figure(
        doc,
        "owner_observed_vs_recovered_demand.png",
        "Hình 4.1: So sánh observed sales và recovered demand",
        "Hinh_4.1_Observed_sales_va_recovered_demand.png",
    )
    add_paragraph(
        doc,
        "Điểm quan trọng là recovered demand không được hiểu như ground truth tuyệt đối. Nó là demand proxy đã được kiểm soát bằng expanding window, calibration và capping. "
        "Do đó, kết quả recovery cần được đọc cùng các kiểm tra uplift, pseudo-stockout validation và sensitivity ở các mục tiếp theo.",
    )

    doc.add_heading("4.1.2. Mức độ Cải thiện từ Cơ chế Điền khuyết", level=3)
    add_paragraph(
        doc,
        f"Cơ chế imputation làm tổng demand tăng {fmt_pct(uplift.get('Recovered lift over observed', ''))} so với observed sales. "
        f"Phần lost demand sau cap giữ lại {fmt_pct(uplift.get('Cap retained share of raw lost demand', ''))} so với raw calibrated lost demand trước cap. "
        "Điều này cho thấy imputation có tác động rõ ràng nhưng không giữ toàn bộ giá trị raw, nhờ vậy giảm rủi ro phóng đại nhu cầu.",
    )
    add_table(
        doc,
        ["Chỉ tiêu", "Giá trị"],
        [
            ["Raw calibrated lost demand before cap", fmt_float(uplift.get("Raw calibrated lost demand before cap", ""), 2)],
            ["Recovered lost demand after cap", fmt_float(uplift.get("Recovered lost demand", ""), 2)],
            ["Cap retained share", fmt_pct(uplift.get("Cap retained share of raw lost demand", ""))],
            ["Mean stockout uplift per stockout row", fmt_float(uplift.get("Mean stockout uplift per stockout row", ""), 4)],
        ],
        "Bảng 4.2: Tác động của imputation lên recovered demand",
    )
    add_figure(
        doc,
        "imputation_daily_uplift.png",
        "Hình 4.2: Uplift sau imputation theo ngày",
        "Hinh_4.2_Uplift_sau_imputation_theo_ngay.png",
    )
    add_figure(
        doc,
        "imputation_series_lift_distribution.png",
        "Hình 4.3: Phân phối uplift theo chuỗi",
        "Hinh_4.3_Phan_phoi_uplift_theo_chuoi.png",
    )
    add_paragraph(
        doc,
        "Ngoài lượng uplift tổng, đồ án kiểm tra thêm hai tín hiệu có ý nghĩa nghiệp vụ trong các giờ stockout. "
        "Tín hiệu velocity/momentum cho biết sản phẩm có đang bán nhanh hơn nền lịch sử ngắn hạn hay không trước khi mất quan sát. "
        "Tín hiệu peer/substitution cho biết các sản phẩm cùng cửa hàng và cùng nhóm hàng có đang tăng sales hoặc cùng gặp stockout hay không. "
        "Hai kiểm tra này giúp tránh cách hiểu quá đơn giản rằng mọi zero hoặc giảm sales đều được impute như nhau; recovery được đặt trong bối cảnh bán hàng trước stockout và bối cảnh sản phẩm thay thế.",
    )
    if sub_velocity_rows:
        add_table(
            doc,
            [
                "Segment",
                "Rows",
                "Velocity 3h/24h",
                "Momentum 3h-6h",
                "Peer sales",
                "Peer velocity 3h/24h",
                "Peer stockout rate",
            ],
            sub_velocity_rows,
            "Bảng 4.3: Tín hiệu velocity và peer/substitution tại các giờ stockout",
        )
        add_figure(
            doc,
            "stockout_substitution_velocity_diagnostics.png",
            "Hình 4.4: So sánh tín hiệu velocity và peer giữa stockout và non-stockout",
            "Hinh_4.4_Tin_hieu_velocity_peer_stockout.png",
        )

    doc.add_heading("4.1.3. Kiểm định Dựa trên Hiện tượng Hết hàng Giả định", level=3)
    add_paragraph(
        doc,
        "Pseudo-stockout validation được dùng để kiểm tra chất lượng recovery trong điều kiện có thể quan sát được đáp án gần đúng. "
        "Một phần quan sát non-stockout được giả lập như stockout, sau đó mô hình thử khôi phục và so sánh với giá trị thực đã biết. "
        f"Kết quả ở cấp aggregate cho thấy WAPE daily aggregate đạt {fmt_pct(pseudo_daily_wape)}, còn hourly aggregate đạt {fmt_pct(pseudo_hourly_wape)}.",
    )
    add_table(
        doc,
        ["Validation level", "RMSE", "MAE", "WAPE", "sMAPE", "Prediction / actual ratio", "N"],
        df_rows(
            "imputation_pseudo_stockout_aggregate_validation.csv",
            ["Validation level", "RMSE", "MAE", "WAPE", "sMAPE", "Prediction / actual ratio", "N"],
        ),
        "Bảng 4.4: Kết quả pseudo-stockout validation",
    )
    add_paragraph(
        doc,
        f"Prediction/actual ratio ở cấp daily aggregate là {fmt_float(pseudo_ratio, 4)}. "
        "Chỉ số này giúp kiểm tra recovery có xu hướng dự đoán cao hoặc thấp hơn giá trị thực trong bài kiểm tra giả lập hay không. "
        "Đây là lý do calibration và cap q90 là cần thiết: nếu không kiểm soát, imputation có thể làm recovered demand tăng quá mức.",
    )

    doc.add_heading("4.1.4. Phân tích Độ nhạy theo Ngưỡng Giới hạn", level=3)
    add_paragraph(
        doc,
        "Phân tích sensitivity kiểm tra mức độ phụ thuộc của recovered demand vào ngưỡng cap lost demand. "
        f"Khi cap ở q50, recovered lift chỉ đạt {fmt_pct(q50_lift)}, khá bảo thủ. Khi dùng q100, recovered lift tăng lên {fmt_pct(q100_lift)}, nhưng giữ toàn bộ phần lost demand raw và có nguy cơ quá lạc quan. "
        f"Ngưỡng q90 được chọn vì tạo mức lift trung gian {fmt_pct(q90_lift)}, giữ lại khoảng {fmt_pct(q90_share)} lost demand raw.",
    )
    add_table(
        doc,
        ["Scenario", "Cap value", "Recovered demand", "Recovered lost demand", "Recovered lift over observed", "Share of original recovered lost demand"],
        df_rows(
            "imputation_cap_sensitivity.csv",
            ["Scenario", "Cap value", "Recovered demand", "Recovered lost demand", "Recovered lift over observed", "Share of original recovered lost demand"],
        ),
        "Bảng 4.5: Sensitivity analysis theo ngưỡng cap",
    )
    add_figure(
        doc,
        "imputation_top_series_lift.png",
        "Hình 4.5: Các chuỗi có uplift cao sau imputation",
        "Hinh_4.5_Cac_chuoi_co_uplift_cao.png",
    )

    doc.add_heading("4.2. Đánh giá Hiệu suất Dự báo", level=2)
    doc.add_heading("4.2.1. Đối chiếu Mô hình Không Điền khuyết và Có Điền khuyết", level=3)
    add_paragraph(
        doc,
        "Khi đánh giá trên recovered latent demand proxy, mô hình forecast trực tiếp observed sales đạt WAPE "
        f"{fmt_pct(observed_wape)}, trong khi mô hình forecast trên recovered demand đạt WAPE {fmt_pct(recovered_wape)}. "
        "Điều này cho thấy target sau recovery giúp giảm sai số so với việc học trực tiếp từ observed sales bị kiểm duyệt.",
    )
    add_table(
        doc,
        ["Model", "RMSE", "MAE", "WAPE", "sMAPE", "WPE"],
        selected_forecast_rows(),
        "Bảng 4.6: So sánh các mô hình trên recovered latent demand proxy",
    )
    add_figure(
        doc,
        "owner_two_stage_forecast_comparison.png",
        "Hình 4.6: So sánh WAPE giữa các mô hình forecasting",
        "Hinh_4.6_So_sanh_WAPE_giua_cac_mo_hinh.png",
    )
    add_paragraph(
        doc,
        f"Khoảng cách giữa observed-sales forecasting và recovered-demand forecasting cho thấy lợi ích của việc xử lý stockout trước khi forecast. "
        f"Tuy nhiên, recovered-demand LightGBM chưa vượt rõ seasonal naive, chứng tỏ pattern tuần là tín hiệu rất mạnh trong dữ liệu.",
    )

    doc.add_heading("4.2.2. Seasonal Naive với Vai trò Baseline Trọng yếu", level=3)
    add_paragraph(
        doc,
        f"Seasonal naive 7-day đạt WAPE {fmt_pct(seasonal_wape)}, tốt hơn nhiều so với naive x7 và rolling mean 14-day. "
        "Kết quả này nhất quán với phân tích seasonality ở Chương 2: nhu cầu bán lẻ có cấu trúc lặp lại theo tuần, nên dự báo theo cùng kỳ tuần trước là benchmark mạnh.",
    )
    add_paragraph(
        doc,
        "Điểm đáng chú ý là seasonal naive không cần huấn luyện nhưng vẫn rất cạnh tranh. Vì vậy, nếu một mô hình học máy không vượt được seasonal naive, mô hình đó chưa chứng minh được giá trị thực tiễn. "
        "Đây là lý do Chương 4 không chỉ so với naive đơn giản mà đặt seasonal naive làm baseline trọng yếu.",
    )

    doc.add_heading("4.2.3. Hiệu quả của Hybrid Seasonal-ML", level=3)
    add_paragraph(
        doc,
        f"Hybrid seasonal-ML đạt WAPE {fmt_pct(hybrid_wape)}, là kết quả tốt nhất trong các mô hình được so sánh. "
        f"So với observed-sales forecasting, WAPE giảm tương đối khoảng {fmt_pct(rel_improve)}. "
        "Mô hình hybrid hoạt động tốt vì kết hợp hai nguồn tín hiệu: seasonal naive giữ pattern tuần ổn định, còn LightGBM hiệu chỉnh theo lịch sử demand, stockout, định danh sản phẩm-cửa hàng và các biến lịch.",
    )
    add_figure(
        doc,
        "owner_two_stage_forecast_interval.png",
        "Hình 4.7: Dự báo và khoảng dự báo của mô hình two-stage",
        "Hinh_4.7_Du_bao_va_khoang_du_bao_two_stage.png",
    )
    add_figure(
        doc,
        "owner_two_stage_bias_comparison.png",
        "Hình 4.8: So sánh xu hướng bias giữa các mô hình",
        "Hinh_4.8_So_sanh_bias_giua_cac_mo_hinh.png",
    )

    doc.add_heading("4.2.4. Đóng góp của Đặc trưng", level=3)
    add_paragraph(
        doc,
        "Feature importance của recovered-demand LightGBM cho thấy các biến quan trọng nhất gồm stockout_rate và các lag của recovered demand như lag 28, lag 1, lag 14 và lag 7. "
        "Điều này phù hợp với logic của bài toán: mức độ stockout giải thích phần doanh số bị kiểm duyệt, còn các lag theo tuần và nhiều tuần phản ánh seasonality đã quan sát trong EDA.",
    )
    add_table(
        doc,
        ["Feature", "Importance"],
        df_rows("owner_recovereddemand_forecasting_feature_importance.csv", ["feature", "importance"], max_rows=12),
        "Bảng 4.7: Top đặc trưng quan trọng của recovered-demand LightGBM",
    )
    add_figure(
        doc,
        "owner_recovereddemand_forecasting_feature_importance_top20.png",
        "Hình 4.9: Top 20 feature importance của recovered-demand LightGBM",
        "Hinh_4.9_Top20_feature_importance.png",
    )

    doc.add_heading("4.3. Phân tích Diagnostics và Độ tin cậy của Kết quả", level=2)
    add_paragraph(
        doc,
        "Ngoài WAPE, diagnostics được sử dụng để kiểm tra phần dư và xu hướng sai lệch của mô hình. "
        "Bảng diagnostics cho thấy hybrid có WAPE thấp nhất trong nhóm so sánh chính, nhưng phần dư vẫn chưa hoàn toàn là white noise theo mọi kiểm định. "
        "Điều này là bình thường với dữ liệu bán lẻ nhiều zero, nhiều chuỗi và chịu ảnh hưởng stockout; mục tiêu thực tế là giảm sai số vận hành chứ không giả định phần dư hoàn toàn chuẩn.",
    )
    add_table(
        doc,
        ["Model", "WAPE", "WPE", "Residual mean", "Ljung-Box p-value", "N"],
        df_rows(
            "owner_two_stage_diagnostics.csv",
            ["Model", "WAPE", "WPE", "Residual mean", "Ljung-Box p-value", "N"],
        ),
        "Bảng 4.8: Diagnostics phần dư của các mô hình",
    )
    add_figure(
        doc,
        "owner_two_stage_residual_acf.png",
        "Hình 4.10: ACF phần dư của mô hình two-stage",
        "Hinh_4.10_ACF_phan_du_two_stage.png",
    )
    add_figure(
        doc,
        "owner_two_stage_residual_distribution.png",
        "Hình 4.11: Phân phối phần dư của mô hình two-stage",
        "Hinh_4.11_Phan_phoi_phan_du_two_stage.png",
    )
    add_paragraph(
        doc,
        "Kết quả non-overlapping 7-day targets cũng được dùng để kiểm tra độ ổn định khi các target window không chồng lấn. "
        "Hybrid vẫn giữ WAPE thấp nhất trong nhóm so sánh, cho thấy kết luận không chỉ đến từ cách tạo nhiều origin chồng lấn trong tập test.",
    )
    add_table(
        doc,
        ["Model", "Evaluation mode", "WAPE", "WPE", "Residual mean", "N"],
        df_rows(
            "owner_two_stage_nonoverlap_diagnostics.csv",
            ["Model", "Evaluation mode", "WAPE", "WPE", "Residual mean", "N"],
        ),
        "Bảng 4.9: Diagnostics trên các target window không chồng lấn",
    )

    doc.add_heading("Kết luận Chương 4", level=2)
    add_paragraph(
        doc,
        "Kết quả thực nghiệm cho thấy hướng tiếp cận two-stage là hợp lý với dữ liệu bán lẻ có stockout. "
        "Recovery giúp tăng demand proxy khoảng "
        f"{fmt_pct(recovery.get('Recovered lift over observed', ''))}, phản ánh phần nhu cầu có thể bị che khuất trong observed sales. "
        "Các kiểm tra pseudo-stockout và sensitivity cho thấy imputation có tác động đáng kể nhưng cần được calibration và cap để tránh phóng đại.",
    )
    add_paragraph(
        doc,
        f"Ở bài toán forecasting, seasonal naive 7-day là baseline rất mạnh với WAPE {fmt_pct(seasonal_wape)}. "
        f"Mô hình tốt nhất là hybrid seasonal-ML với WAPE {fmt_pct(hybrid_wape)}, cải thiện so với observed-sales forecasting và nhỉnh hơn seasonal naive. "
        "Kết quả này củng cố insight chính của đồ án: stockout cần được xử lý trước khi forecast, nhưng seasonality tuần vẫn là tín hiệu nền tảng không nên bỏ qua.",
    )

    saved_path = save_doc(doc, OUT_PATH)
    print(saved_path)


if __name__ == "__main__":
    build_doc()

