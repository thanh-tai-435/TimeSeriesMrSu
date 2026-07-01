from __future__ import annotations

import csv
import shutil
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.dates as mdates
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TABLE_DIR = BASE_DIR / "outputs" / "tables"
FIGURE_DIR = BASE_DIR / "outputs" / "figures"
OUT_DIR = BASE_DIR / "deliverables" / "reports"
OUT_PATH = OUT_DIR / "Phan_mo_dau_va_Chuong_1_Du_lieu_Tien_xu_ly.docx"
IMAGE_OUT_DIR = OUT_DIR / "Phan_mo_dau_Chuong_1_Hinh_anh_goc"


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
        doc.save(fallback)
        return fallback


def set_default_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12 if style_name == "Normal" else 14)
        style.font.bold = style_name.startswith("Heading")
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)

    for style_name, italic, bold in [("Caption", True, False), ("Table Caption", False, True)]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.font.italic = italic
        style.font.bold = bold
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    doc.add_paragraph(caption, style="Table Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, export_filename: str, width: float = 6.2) -> None:
    path = FIGURE_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[Thiếu hình: {filename}]")
        return
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, IMAGE_OUT_DIR / export_filename)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")


def read_metric_table(filename: str) -> dict[str, str]:
    result: dict[str, str] = {}
    with (TABLE_DIR / filename).open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            metric_key = "Metric" if "Metric" in row else "metric"
            value_key = "Value" if "Value" in row else "value"
            result[row[metric_key]] = row[value_key]
    return result


def fmt_float(value: str | float, digits: int = 4) -> str:
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)


def fmt_pct(value: str | float, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def fmt_int(value: str | float) -> str:
    try:
        return f"{int(float(value)):,}"
    except Exception:
        return str(value)


def format_cell(col: str, value) -> str:
    if isinstance(value, float):
        lowered = col.lower()
        if any(key in lowered for key in ["rate", "wape", "wpe", "lift", "share", "fraction"]):
            return fmt_pct(value)
        return fmt_float(value, 4)
    return str(value)


def table_rows(filename: str, columns: list[str]) -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / filename)
    return [[format_cell(col, row[col]) for col in columns] for _, row in df.iterrows()]


def raw_shape_rows() -> list[list[str]]:
    rows = []
    for label, path in [
        ("Raw train", BASE_DIR / "data" / "raw" / "train.parquet"),
        ("Raw eval", BASE_DIR / "data" / "raw" / "eval.parquet"),
        ("Base sau chuẩn hóa", BASE_DIR / "data" / "processed" / "fresh50k_base.parquet"),
        ("Sample daily 10%", BASE_DIR / "data" / "sample" / "fresh50k_sample_010.parquet"),
        ("Sample hourly 10%", BASE_DIR / "data" / "sample" / "fresh50k_hourly_sample_010.parquet"),
    ]:
        df = pd.read_parquet(path)
        rows.append(
            [
                label,
                fmt_int(len(df)),
                fmt_int(df["series_id"].nunique()) if "series_id" in df.columns else "Chưa tạo",
                str(pd.to_datetime(df["dt"]).min()) if "dt" in df.columns else "",
                str(pd.to_datetime(df["dt"]).max()) if "dt" in df.columns else "",
            ]
        )
    return rows


def make_pipeline_figure() -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / "data_preprocessing_pipeline.png"
    steps = [
        ("Raw train/eval\nparquet", 0.08, 0.62),
        ("Chuẩn hóa\ncột/ngày/id", 0.28, 0.62),
        ("Sample 10%\ntheo series", 0.48, 0.62),
        ("Daily → Hourly\nhours_sale/status", 0.68, 0.62),
        ("Feature table\nlag/rolling/stockout", 0.88, 0.62),
        ("Time split\ntrain/val/test", 0.48, 0.26),
        ("Recovery + daily\nforecasting target", 0.76, 0.26),
    ]
    fig, ax = plt.subplots(figsize=(12, 4.5))
    ax.axis("off")
    for text, x, y in steps:
        ax.text(
            x,
            y,
            text,
            ha="center",
            va="center",
            fontsize=10,
            bbox=dict(boxstyle="round,pad=0.35", facecolor="#ecfdf5", edgecolor="#047857", linewidth=1.5),
            transform=ax.transAxes,
        )
    arrows = [
        ((0.16, 0.62), (0.22, 0.62)),
        ((0.36, 0.62), (0.42, 0.62)),
        ((0.56, 0.62), (0.62, 0.62)),
        ((0.76, 0.62), (0.82, 0.62)),
        ((0.48, 0.54), (0.48, 0.34)),
        ((0.56, 0.26), (0.68, 0.26)),
    ]
    for start, end in arrows:
        ax.annotate("", xy=end, xytext=start, xycoords="axes fraction", arrowprops=dict(arrowstyle="->", lw=1.6, color="#374151"))
    ax.set_title("Quy trình dữ liệu và tiền xử lý", fontsize=13, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def make_split_timeline_figure() -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    split = pd.read_csv(TABLE_DIR / "split_summary.csv")
    split["Start date"] = pd.to_datetime(split["Start date"])
    split["End date"] = pd.to_datetime(split["End date"])
    colors = {"Train": "#059669", "Validation": "#d97706", "Test": "#2563eb"}
    fig, ax = plt.subplots(figsize=(10, 2.6))
    for idx, row in split.iterrows():
        start = row["Start date"]
        end = row["End date"]
        ax.barh(
            y=0,
            width=end - start,
            left=start,
            height=0.34,
            color=colors[row["Split"]],
            label=row["Split"],
            alpha=0.9,
        )
        ax.text(start + (end - start) / 2, 0, row["Split"], ha="center", va="center", color="white", fontsize=10, fontweight="bold")
    ax.set_yticks([])
    ax.xaxis.set_major_formatter(mdates.DateFormatter("%d-%m"))
    ax.set_xlabel("Thời gian")
    ax.set_title("Phân chia Train/Validation/Test theo trục thời gian", fontsize=12, fontweight="bold")
    ax.grid(axis="x", alpha=0.25)
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(FIGURE_DIR / "time_split_timeline.png", dpi=300, bbox_inches="tight")
    plt.close(fig)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if IMAGE_OUT_DIR.exists():
        shutil.rmtree(IMAGE_OUT_DIR)
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_pipeline_figure()
    make_split_timeline_figure()

    quality = read_metric_table("data_quality_summary.csv")
    recovery = read_metric_table("owner_latent_recovery_summary.csv")
    forecast = pd.read_csv(TABLE_DIR / "owner_two_stage_forecasting_comparison.csv")
    forecast = forecast[forecast["Evaluation target"] == "Recovered latent demand proxy"]
    hybrid_wape = float(forecast[forecast["Model"] == "Recovered seasonal-ML hybrid"]["WAPE"].iloc[0])
    observed_wape = float(forecast[forecast["Model"] == "Observed-sales forecasting"]["WAPE"].iloc[0])

    doc = Document()
    set_default_font(doc)

    title = doc.add_heading("PHẦN MỞ ĐẦU VÀ TỔNG QUAN", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("0.1. Bối cảnh và Động lực Nghiên cứu", level=2)
    add_paragraph(
        doc,
        "Trong bán lẻ hàng tươi sống, dự báo nhu cầu không chỉ là bài toán thống kê mà còn là bài toán vận hành. Nếu dự báo thấp hơn nhu cầu thực, cửa hàng có thể hết hàng, mất doanh số và làm giảm trải nghiệm khách hàng. Nếu dự báo quá cao, cửa hàng có thể nhập dư, làm tăng tồn kho, hao hụt và chi phí hủy hàng. Vì vậy, forecast cần phục vụ trực tiếp quyết định nhập hàng, không chỉ tối ưu một chỉ số mô hình.",
    )
    add_paragraph(
        doc,
        "Điểm khó của bộ dữ liệu FreshRetailNet-50K là observed sales không luôn bằng demand. Khi sản phẩm stockout, doanh số ghi nhận bị giới hạn bởi tồn kho. Một thời điểm bán thấp có thể không phản ánh khách hàng ít mua, mà phản ánh cửa hàng không còn hàng để bán. Nếu mô hình học trực tiếp từ observed sales, nó có thể học sai rằng nhu cầu thấp tại các thời điểm hết hàng. Đây là động lực chính để đồ án tiếp cận bài toán theo hướng khôi phục latent demand trước khi forecast.",
    )
    add_paragraph(
        doc,
        "Bên cạnh stockout, dữ liệu còn có nhiều chuỗi cửa hàng-sản phẩm, nhiều giá trị 0, seasonality theo giờ/tuần và khác biệt lớn giữa các nhóm chuỗi. Điều này khiến cách fit riêng một mô hình cổ điển cho từng chuỗi trở nên khó mở rộng. Đồ án vì vậy chọn một framework hai giai đoạn: xử lý stockout ở cấp giờ, sau đó tổng hợp lên cấp ngày để dự báo tổng nhu cầu 7 ngày tiếp theo.",
    )

    doc.add_heading("0.2. Mục tiêu và Câu hỏi Nghiên cứu", level=2)
    add_paragraph(
        doc,
        "Mục tiêu tổng quát của đề tài là xây dựng một pipeline dự báo nhu cầu bán lẻ có xét đến hiện tượng stockout. Pipeline cần vừa đủ chính xác để so sánh với baseline, vừa đủ minh bạch để giải thích với góc nhìn vận hành. Bài toán cuối cùng được đặt ở cấp daily next-7-day demand, vì đây là horizon phù hợp với quyết định replenishment theo tuần.",
    )
    add_table(
        doc,
        ["Câu hỏi nghiên cứu", "Cách trả lời trong đồ án"],
        [
            ["Observed sales có bị sai lệch bởi stockout không?", "Phân tích stockout rate, uplift recovered demand và lost sales proxy"],
            ["Có thể recovery latent demand mà không leakage thời gian không?", "Dùng expanding-window recovery, mỗi block chỉ học từ non-stockout rows trong quá khứ"],
            ["Imputation có làm demand tăng quá mức không?", "Dùng calibration, cap q90, pseudo-stockout validation và sensitivity analysis"],
            ["Forecast trên recovered demand có tốt hơn observed sales không?", "So sánh observed-sales LightGBM, recovered-demand LightGBM, seasonal naive và hybrid"],
            ["Kết quả có ý nghĩa vận hành gì?", "Diễn giải lost sales, seasonality tuần, prediction intervals và quy trình replenishment"],
        ],
        "Bảng 0.1: Mục tiêu và câu hỏi nghiên cứu của đề tài",
    )

    doc.add_heading("0.3. Những Đóng góp Cốt lõi của Đề tài", level=2)
    add_paragraph(
        doc,
        "Đóng góp thứ nhất là framing lại bài toán: thay vì forecast observed sales một cách trực tiếp, đồ án xem observed sales là dữ liệu có thể bị censor bởi stockout. Cách framing này phù hợp hơn với mục tiêu vận hành, vì replenishment cần dự báo nhu cầu tiềm năng chứ không chỉ lượng hàng đã bán.",
    )
    add_paragraph(
        doc,
        "Đóng góp thứ hai là xây dựng quy trình latent demand recovery ở cấp giờ bằng expanding window. Trong train period, nhu cầu của mỗi block chỉ được recovery bằng mô hình học từ dữ liệu non-stockout trong quá khứ. Cơ chế này tránh việc dùng dữ liệu tương lai để khôi phục quá khứ, một lỗi leakage thường gặp khi xử lý time series.",
    )
    add_paragraph(
        doc,
        "Đóng góp thứ ba là kiểm soát imputation bằng nhiều lớp: calibration, capping, pseudo-stockout validation và sensitivity analysis. Nhờ đó, recovered demand được xem là proxy có kiểm soát chứ không bị trình bày như ground truth tuyệt đối.",
    )
    add_paragraph(
        doc,
        "Đóng góp thứ tư là kết hợp seasonal naive và LightGBM trong hybrid seasonal-ML. Kết quả EDA cho thấy seasonality tuần rất mạnh, nên seasonal naive là baseline quan trọng. LightGBM giúp hiệu chỉnh theo recovered demand, stockout, định danh cửa hàng-sản phẩm và các đặc trưng lịch sử. Hybrid vì vậy tận dụng cả quy luật mùa vụ ổn định và khả năng học phi tuyến của mô hình ML.",
    )
    add_figure(
        doc,
        "owner_expanding_window_process.png",
        "Hình 0.1: Khung tiếp cận tổng quát của đề tài",
        "Hinh_0.1_Khung_tiep_can_tong_quat.png",
    )

    doc.add_heading("0.4. Phạm vi Nghiên cứu và Các Giới hạn Chuyên môn (Disclaimer/Out of Scope)", level=2)
    add_paragraph(
        doc,
        "Đồ án sử dụng sample 10% theo series_id để phù hợp tài nguyên tính toán. Cách lấy mẫu này giữ nguyên toàn bộ trục thời gian của các chuỗi được chọn, nhưng không đại diện đầy đủ cho mọi quan hệ giữa 50.000 chuỗi trong dữ liệu gốc. Vì vậy, kết quả nên được hiểu là kết quả thực nghiệm trên sample có kiểm soát, không phải benchmark cuối cùng trên toàn bộ dữ liệu.",
    )
    add_paragraph(
        doc,
        "Recovered demand trong đồ án là proxy, không phải demand thật được quan sát trực tiếp. Khi stockout xảy ra, không có dữ liệu ground truth tuyệt đối về khách hàng muốn mua bao nhiêu. Vì vậy, lost sales được báo cáo như một ước lượng có kiểm soát, cần đọc cùng pseudo-stockout validation và sensitivity analysis.",
    )
    add_paragraph(
        doc,
        "Đồ án không triển khai các mô hình deep imputation nặng như SAITS, CSDI, TimesNet hay ImputeFormer trong pipeline chính. Lý do là mục tiêu môn học ưu tiên một pipeline chặt chẽ, có thể tái lập, có kiểm soát leakage và có diễn giải rõ. Các mô hình deep learning được xem là hướng phát triển tương lai khi có thêm thời gian, tài nguyên và thiết kế validation đủ nghiêm ngặt.",
    )
    add_table(
        doc,
        ["Trong phạm vi", "Ngoài phạm vi"],
        [
            ["Sample 10% theo series, giữ nguyên timeline", "Benchmark full 50K series trên toàn bộ dữ liệu"],
            ["Hourly recovery và daily next-7-day forecasting", "Forecast từng giờ cho replenishment thời gian thực"],
            ["LightGBM, seasonal naive và hybrid", "Deep imputation/deep forecasting quy mô lớn"],
            ["Pseudo-stockout validation và sensitivity", "Xác nhận demand thật bằng giao dịch khách hàng không quan sát được"],
            ["Business insight định hướng vận hành", "Tối ưu hóa tồn kho đầy đủ có chi phí, lead time và ràng buộc cung ứng thực tế"],
        ],
        "Bảng 0.2: Phạm vi và giới hạn của đề tài",
    )

    doc.add_heading("0.5. Tổng quan các Kết quả Đạt được", level=2)
    add_paragraph(
        doc,
        "Trên sample 10%, dữ liệu hourly gồm "
        f"{fmt_int(quality.get('Number of rows', ''))} quan sát, {fmt_int(quality.get('Number of series', ''))} chuỗi và stockout rate khoảng {fmt_pct(quality.get('Stockout rate', ''))}. "
        "Sau recovery, recovered latent demand cao hơn observed sales khoảng "
        f"{fmt_pct(recovery.get('Recovered lift over observed', ''))}, tương ứng recovered lost demand "
        f"{fmt_float(recovery.get('Recovered lost demand', ''), 2)}.",
    )
    add_paragraph(
        doc,
        "Ở bước forecasting, observed-sales LightGBM đạt WAPE khoảng "
        f"{fmt_pct(observed_wape)}, trong khi hybrid seasonal-ML đạt WAPE khoảng {fmt_pct(hybrid_wape)} trên recovered latent demand proxy. "
        "Kết quả này cho thấy xử lý stockout trước khi forecast giúp giảm sai số so với học trực tiếp từ observed sales bị censor.",
    )
    add_table(
        doc,
        ["Kết quả", "Giá trị"],
        [
            ["Hourly sample rows", fmt_int(quality.get("Number of rows", ""))],
            ["Number of series", fmt_int(quality.get("Number of series", ""))],
            ["Stockout rate", fmt_pct(quality.get("Stockout rate", ""))],
            ["Recovered lift over observed", fmt_pct(recovery.get("Recovered lift over observed", ""))],
            ["Observed-sales forecasting WAPE", fmt_pct(observed_wape)],
            ["Hybrid seasonal-ML WAPE", fmt_pct(hybrid_wape)],
        ],
        "Bảng 0.3: Tóm tắt các kết quả chính của đề tài",
    )

    doc.add_heading("0.6. Tổng quan Tài liệu và Cơ sở Lý thuyết (Literature Review)", level=2)
    add_paragraph(
        doc,
        "Bộ dữ liệu FreshRetailNet-50K được giới thiệu như một benchmark bán lẻ có nhiều chuỗi cửa hàng-sản phẩm, chứa thông tin bán hàng theo ngày, chuỗi doanh số theo giờ, trạng thái tồn kho theo giờ và các biến ngoại sinh như khuyến mãi/thời tiết [Nguồn: FreshRetailNet-50K paper/repository, arXiv:2505.16319]. Đặc điểm quan trọng của bộ dữ liệu là observed sales có thể bị censor bởi stockout, vì vậy bài toán demand forecasting gắn chặt với latent demand recovery.",
    )
    add_paragraph(
        doc,
        "Trong lý thuyết chuỗi thời gian, stationarity, ACF/PACF và seasonality là các công cụ nền tảng để nhận diện cấu trúc phụ thuộc theo thời gian. Tuy nhiên, với dữ liệu bán lẻ nhiều chuỗi, nhiều zero và nhiều covariates, việc áp dụng mô hình cổ điển riêng cho từng chuỗi thường khó mở rộng. Do đó, đồ án sử dụng các công cụ này như bước chẩn đoán và thiết kế đặc trưng, thay vì ép toàn bộ pipeline theo ARIMA/SARIMA.",
    )
    add_paragraph(
        doc,
        "LightGBM và các mô hình gradient boosting thường phù hợp với dữ liệu bảng có nhiều đặc trưng lag, rolling, calendar, định danh và biến ngoại sinh [Nguồn: LightGBM original paper/documentation]. Trong đồ án, LightGBM được dùng như global model để học pattern chung từ nhiều chuỗi, còn seasonal naive được giữ làm benchmark mùa vụ mạnh. Hybrid seasonal-ML là cách kết hợp hai hướng: giữ nhịp tuần ổn định và cho phép mô hình học máy hiệu chỉnh phần phi tuyến.",
    )

    title = doc.add_heading("CHƯƠNG 1: NGUỒN DỮ LIỆU VÀ TIỀN XỬ LÝ", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1.1. Tổng quan Cấu trúc Nguồn Dữ liệu", level=2)
    add_paragraph(
        doc,
        "Nguồn dữ liệu ban đầu gồm hai file parquet: train.parquet và eval.parquet. Sau khi ghép và chuẩn hóa, dữ liệu base có 4.850.000 dòng daily, tương ứng 50.000 chuỗi cửa hàng-sản phẩm trong giai đoạn từ 28/03/2024 đến 02/07/2024. Mỗi dòng daily chứa tổng sale_amount trong ngày, mảng hours_sale gồm 24 giá trị theo giờ và mảng hours_stock_status mô tả trạng thái stockout theo giờ.",
    )
    add_table(
        doc,
        ["Tập dữ liệu", "Số dòng", "Số chuỗi", "Bắt đầu", "Kết thúc"],
        raw_shape_rows(),
        "Bảng 1.1: Kích thước dữ liệu qua các bước xử lý",
    )
    add_paragraph(
        doc,
        "Các cột định danh chính gồm city_id, store_id, product_id và các cấp category như management_group_id, first_category_id, second_category_id, third_category_id. Các biến ngoại sinh gồm discount, holiday_flag, activity_flag và các biến thời tiết như precip, temp, humidity, wind. Những biến này được giữ lại để phục vụ feature engineering ở các chương sau.",
    )

    doc.add_heading("1.2. Thiết lập Định danh Chuỗi Thời gian (Time-series Identifiers)", level=2)
    add_paragraph(
        doc,
        "Đơn vị chuỗi thời gian của đồ án là một cặp cửa hàng-sản phẩm trong một thành phố. Vì vậy, series_id được tạo bằng cách ghép city_id, store_id và product_id. Cách định danh này bảo đảm mỗi chuỗi biểu diễn lịch sử bán hàng của cùng một sản phẩm tại cùng một cửa hàng, tránh trộn lẫn nhu cầu giữa các cửa hàng hoặc sản phẩm khác nhau.",
    )
    add_table(
        doc,
        ["Thành phần", "Vai trò"],
        [
            ["city_id", "Phân biệt thành phố/khu vực vận hành"],
            ["store_id", "Phân biệt cửa hàng"],
            ["product_id", "Phân biệt sản phẩm"],
            ["series_id = city_id_store_id_product_id", "Định danh duy nhất cho một chuỗi cửa hàng-sản phẩm"],
            ["category ids", "Bổ sung ngữ cảnh nhóm hàng và hỗ trợ peer/substitution features"],
        ],
        "Bảng 1.2: Cấu trúc định danh chuỗi thời gian",
    )
    add_paragraph(
        doc,
        "Việc định danh chuỗi trước khi sampling và feature engineering là rất quan trọng. Nếu lấy mẫu theo từng dòng ngẫu nhiên, timeline của chuỗi sẽ bị đứt đoạn và làm sai các đặc trưng lag/rolling. Ngược lại, lấy mẫu theo series_id giữ nguyên toàn bộ lịch sử của chuỗi được chọn.",
    )

    doc.add_heading("1.3. Phân rã Dữ liệu (Disaggregation) từ Cấp độ Ngày sang Giờ", level=2)
    add_paragraph(
        doc,
        "Dữ liệu gốc ở cấp ngày nhưng chứa hai mảng 24 phần tử: hours_sale và hours_stock_status. Để xử lý stockout đúng bản chất, đồ án phân rã mỗi dòng daily thành 24 dòng hourly. Sau phân rã, cột dt được cộng thêm số giờ tương ứng, sale_amount trở thành doanh số của từng giờ, stockout_flag biểu diễn trạng thái stockout của từng giờ, còn daily_sale_amount được giữ như thông tin tổng ngày tham chiếu.",
    )
    add_figure(
        doc,
        "data_preprocessing_pipeline.png",
        "Hình 1.1: Quy trình dữ liệu và tiền xử lý",
        "Hinh_1.1_Quy_trinh_du_lieu_va_tien_xu_ly.png",
    )
    add_paragraph(
        doc,
        "Việc phân rã hourly là cần thiết vì stockout xảy ra theo giờ. Nếu chỉ làm việc ở cấp ngày, một ngày có vài giờ hết hàng và vài giờ còn hàng sẽ bị gom thành một quan sát duy nhất, làm mất thông tin quan trọng về thời điểm demand bị censor. Recovery ở cấp giờ giúp mô hình ước lượng lost demand chính xác hơn trước khi tổng hợp lại thành daily demand cho forecasting.",
    )

    doc.add_heading("1.4. Phương pháp Lấy mẫu (Sampling) theo Chuỗi Dữ liệu", level=2)
    add_paragraph(
        doc,
        "Do giới hạn tài nguyên tính toán, đồ án sử dụng sample_frac = 0.1, tương ứng 10% số chuỗi. Việc lấy mẫu được thực hiện theo series_id, không lấy mẫu ngẫu nhiên theo dòng. Với 50.000 chuỗi ban đầu, sample 10% giữ lại 5.000 chuỗi đầy đủ lịch sử thời gian. Sau khi phân rã hourly, sample có 11.640.000 quan sát.",
    )
    add_table(
        doc,
        ["Chỉ tiêu", "Giá trị"],
        [
            ["Sample fraction", fmt_pct(quality.get("Sample fraction", ""))],
            ["Number of hourly rows", fmt_int(quality.get("Number of rows", ""))],
            ["Number of series", fmt_int(quality.get("Number of series", ""))],
            ["Number of stores", fmt_int(quality.get("Number of stores", ""))],
            ["Number of products", fmt_int(quality.get("Number of products", ""))],
            ["Number of cities", fmt_int(quality.get("Number of cities", ""))],
            ["Start date", quality.get("Start date", "")],
            ["End date", quality.get("End date", "")],
        ],
        "Bảng 1.3: Thống kê dữ liệu sau lấy mẫu và phân rã hourly",
    )
    add_paragraph(
        doc,
        "Lấy mẫu theo chuỗi giúp các đặc trưng time series như lag 1 giờ, lag 24 giờ, lag 168 giờ, rolling mean và rolling stockout không bị sai do thiếu đoạn thời gian. Đây là lựa chọn quan trọng để bảo đảm tính đúng đắn kỹ thuật của pipeline trong điều kiện không dùng full data.",
    )

    doc.add_heading("1.5. Phân chia Tập Dữ liệu (Train/Validation/Test Split) theo Trục Thời gian", level=2)
    add_paragraph(
        doc,
        "Vì đây là bài toán chuỗi thời gian, dữ liệu được chia theo trục thời gian thay vì random split. Validation gồm 7 ngày và test gồm 14 ngày cuối, phần còn lại dùng cho train. Cách chia này mô phỏng điều kiện triển khai thực tế: tại thời điểm dự báo chỉ được dùng dữ liệu quá khứ, không được nhìn vào tương lai.",
    )
    add_table(
        doc,
        ["Split", "Start date", "End date", "Rows", "Series"],
        table_rows("split_summary.csv", ["Split", "Start date", "End date", "Rows", "Series"]),
        "Bảng 1.4: Phân chia train/validation/test ở cấp hourly",
    )
    add_figure(
        doc,
        "time_split_timeline.png",
        "Hình 1.2: Timeline phân chia train/validation/test",
        "Hinh_1.2_Timeline_train_validation_test.png",
    )
    add_paragraph(
        doc,
        "Riêng với bài toán daily next-7-day forecasting, split được định nghĩa theo forecast origin. Mỗi origin tại ngày t có target là tổng recovered demand từ t+1 đến t+7. Vì target dùng tương lai 7 ngày, origin cuối cùng của train phải kết thúc sớm hơn split hourly để tránh target window chạm vào validation/test ngoài ý muốn.",
    )
    add_table(
        doc,
        ["Split", "Origin start", "Origin end", "Target window start", "Target window end", "Rows", "Series"],
        table_rows(
            "owner_daily_split_summary.csv",
            ["Split", "Origin start", "Origin end", "Target window start", "Target window end", "Rows", "Series"],
        ),
        "Bảng 1.5: Phân chia dữ liệu cho daily next-7-day forecasting",
    )

    doc.add_heading("1.6. Xử lý Dữ liệu Bất thường và Nhiễu (Outlier/Anomaly Handling)", level=2)
    add_paragraph(
        doc,
        "Kiểm tra chất lượng dữ liệu cho thấy không có missing dt, không có duplicate rows theo series_id-dt và không có missing sale_amount trong sample hourly. Đây là điều kiện quan trọng để tạo lag/rolling features mà không phát sinh lỗ hổng thời gian ngoài ý muốn.",
    )
    add_table(
        doc,
        ["Chỉ tiêu chất lượng", "Giá trị"],
        [
            ["Missing dt count", fmt_int(quality.get("Missing dt count", ""))],
            ["Duplicate rows", fmt_int(quality.get("Duplicate rows", ""))],
            ["Missing sale_amount", fmt_int(quality.get("Missing sale_amount", ""))],
            ["Stockout rate", fmt_pct(quality.get("Stockout rate", ""))],
            ["Average sale_amount", fmt_float(quality.get("Average sale_amount", ""), 4)],
            ["Median sale_amount", fmt_float(quality.get("Median sale_amount", ""), 4)],
        ],
        "Bảng 1.6: Kiểm tra chất lượng dữ liệu sau tiền xử lý",
    )
    add_paragraph(
        doc,
        "Đối với outlier, đồ án không cắt bỏ spike doanh số một cách cơ học. Trong bán lẻ, spike có thể đến từ khuyến mãi, thời tiết, ngày lễ hoặc thay đổi nhu cầu thật; nếu xóa tùy tiện, mô hình sẽ mất tín hiệu vận hành quan trọng. Thay vào đó, pipeline giữ lại các biến ngoại sinh, dùng lag/rolling features, đánh giá bằng WAPE/MAE/RMSE và kiểm soát riêng phần stockout thông qua recovery.",
    )
    add_paragraph(
        doc,
        "Các giá trị 0 cũng không được xem mặc định là lỗi. Một giá trị 0 có thể là không phát sinh nhu cầu, hoặc là kết quả của stockout. Vì vậy, đồ án phân biệt observed sales và stockout_flag. Khi stockout_flag = 1, observed sales được xem là có khả năng bị censor và sẽ được xử lý ở bước latent demand recovery. Khi stockout_flag = 0, giá trị 0 được giữ như một quan sát hợp lệ của demand thấp hoặc không có giao dịch.",
    )
    add_paragraph(
        doc,
        "Tóm lại, tiền xử lý trong Chương 1 hướng tới bảo toàn cấu trúc thời gian và tín hiệu vận hành thay vì làm sạch dữ liệu quá mức. Các bước xử lý chính gồm chuẩn hóa định danh, phân rã hourly, lấy mẫu theo chuỗi, split theo thời gian và gắn nhãn stockout. Đây là nền tảng cho EDA ở Chương 2 và phương pháp two-stage ở Chương 3.",
    )

    saved_path = save_doc(doc, OUT_PATH)
    print(saved_path)


if __name__ == "__main__":
    build_doc()
