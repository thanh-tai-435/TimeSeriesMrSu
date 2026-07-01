from __future__ import annotations

import csv
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TABLE_DIR = BASE_DIR / "outputs" / "tables"
FIGURE_DIR = BASE_DIR / "outputs" / "figures"
OUT_DIR = BASE_DIR / "deliverables" / "reports"
OUT_PATH = OUT_DIR / "Chuong_3_Phuong_phap_nghien_cuu_va_mo_hinh_hoa.docx"
IMAGE_OUT_DIR = OUT_DIR / "Chuong_3_Hinh_anh_goc"


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
        doc.save(fallback)
        return fallback


def set_default_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12 if style_name == "Normal" else 14)
        style.font.bold = style_name.startswith("Heading")
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)

    for style_name, italic, bold in [("Caption", True, False), ("Table Caption", False, True)]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.font.italic = italic
        style.font.bold = bold
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    doc.add_paragraph(caption, style="Table Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, export_filename: str, width: float = 6.2) -> None:
    path = FIGURE_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[Thiếu hình: {filename}]")
        return
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, IMAGE_OUT_DIR / export_filename)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")


def read_metric_table(filename: str) -> dict[str, str]:
    path = TABLE_DIR / filename
    result: dict[str, str] = {}
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            metric_key = "Metric" if "Metric" in row else "metric"
            value_key = "Value" if "Value" in row else "value"
            result[row[metric_key]] = row[value_key]
    return result


def fmt_float(value: str | float, digits: int = 4) -> str:
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)


def fmt_pct(value: str | float, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def df_rows(filename: str, columns: list[str], max_rows: int | None = None) -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / filename)
    if max_rows is not None:
        df = df.head(max_rows)
    rows: list[list[str]] = []
    for _, row in df.iterrows():
        rendered = []
        for col in columns:
            value = row[col]
            if isinstance(value, float):
                if "WAPE" in col or "sMAPE" in col or "WPE" in col or "lift" in col.lower() or "share" in col.lower():
                    rendered.append(fmt_pct(value))
                else:
                    rendered.append(fmt_float(value, 4))
            else:
                rendered.append(str(value))
        rows.append(rendered)
    return rows


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if IMAGE_OUT_DIR.exists():
        shutil.rmtree(IMAGE_OUT_DIR)
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)

    recovery = read_metric_table("owner_latent_recovery_summary.csv")
    hybrid_selection = pd.read_csv(TABLE_DIR / "owner_hybrid_blend_selection.csv").sort_values(
        "Validation WAPE"
    ).iloc[0]

    doc = Document()
    set_default_font(doc)

    title = doc.add_heading("CHƯƠNG 3: PHƯƠNG PHÁP NGHIÊN CỨU VÀ MÔ HÌNH HÓA", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        "Chương 3 trình bày phương pháp nghiên cứu được sử dụng để chuyển từ dữ liệu bán hàng quan sát được sang bài toán dự báo nhu cầu. "
        "Khác với cách dự báo trực tiếp trên doanh số ghi nhận, pipeline của đồ án xem observed sales là một biến có thể bị kiểm duyệt bởi stockout. "
        "Vì vậy, phương pháp được thiết kế theo hai giai đoạn: khôi phục phần nhu cầu bị che khuất ở cấp giờ, sau đó tổng hợp lên cấp ngày để dự báo tổng nhu cầu 7 ngày tiếp theo.",
    )
    add_paragraph(
        doc,
        "So với mục lục ban đầu, chương này bổ sung rõ hai nội dung quan trọng: chia tập theo thời gian để tránh leakage và kiểm soát chất lượng imputation bằng calibration, capping và sensitivity analysis. "
        "Hai nội dung này giúp phương pháp không chỉ có kết quả tốt hơn, mà còn có thể giải thích được về mặt thực nghiệm.",
    )

    doc.add_heading("3.1. Hạn chế của Phương pháp Dự báo Trực tiếp trên Doanh số Quan sát", level=2)
    add_paragraph(
        doc,
        "Trong bối cảnh bán lẻ có stockout, doanh số quan sát được không nhất thiết bằng nhu cầu thực tế. Nếu sản phẩm hết hàng trong một số giờ, lượng bán ghi nhận có thể thấp dù khách hàng vẫn có nhu cầu mua. "
        "Khi dùng trực tiếp observed sales làm target, mô hình có nguy cơ học rằng nhu cầu giảm tại các thời điểm stockout, trong khi nguyên nhân thực tế là ràng buộc tồn kho.",
    )
    add_paragraph(
        doc,
        "Hạn chế này đặc biệt nghiêm trọng với dữ liệu tươi sống vì quyết định bổ sung hàng thường dựa trên nhu cầu kỳ vọng, không chỉ dựa trên số lượng đã bán. "
        "Do đó, dự báo trực tiếp observed sales phù hợp nếu mục tiêu là mô phỏng doanh số đã ghi nhận, nhưng chưa đủ nếu mục tiêu là hỗ trợ vận hành tồn kho và replenishment.",
    )

    doc.add_heading("3.2. Thiết kế Chia tập theo Thời gian và Kiểm soát Rò rỉ Dữ liệu", level=2)
    add_paragraph(
        doc,
        "Toàn bộ pipeline sử dụng split theo thời gian thay vì random split. Lý do là dữ liệu chuỗi thời gian có thứ tự trước-sau rõ ràng; nếu trộn ngẫu nhiên, mô hình có thể nhìn thấy thông tin từ tương lai trong quá trình huấn luyện. "
        "Ở cấp hourly, validation gồm 7 ngày và test gồm 14 ngày cuối. Ở cấp daily forecasting, split được định nghĩa theo forecast origin và target window tương ứng.",
    )
    add_table(
        doc,
        ["Split", "Origin start", "Origin end", "Target window start", "Target window end", "Rows", "Series"],
        df_rows("owner_daily_split_summary.csv", ["Split", "Origin start", "Origin end", "Target window start", "Target window end", "Rows", "Series"]),
        "Bảng 3.1: Chia tập daily forecasting theo forecast origin",
    )
    add_paragraph(
        doc,
        "Thiết kế này đảm bảo mọi bước huấn luyện, hiệu chỉnh và lựa chọn mô hình chỉ sử dụng thông tin có sẵn tại thời điểm dự báo. "
        "Validation được dùng để chọn tham số như trọng số hybrid, còn test chỉ dùng để báo cáo hiệu suất cuối cùng.",
    )

    doc.add_heading("3.3. Khung Phân tích Hai Giai đoạn", level=2)
    add_paragraph(
        doc,
        "Khung hai giai đoạn tách bài toán thành hai nhiệm vụ có bản chất khác nhau. Giai đoạn 1 xử lý stockout ở cấp hourly, vì trạng thái hết hàng xuất hiện theo từng giờ. "
        "Giai đoạn 2 tổng hợp hourly recovered demand thành daily demand và dự báo tổng nhu cầu 7 ngày tiếp theo, phù hợp hơn với quyết định vận hành ở cấp ngày.",
    )
    add_figure(
        doc,
        "owner_expanding_window_process.png",
        "Hình 3.1: Quy trình tổng quát của khung two-stage",
        "Hinh_3.1_Quy_trinh_two_stage.png",
    )
    add_paragraph(
        doc,
        "Cách tiếp cận này giữ được chi tiết cần thiết cho stockout recovery nhưng không buộc mô hình forecasting phải dự báo ở cấp giờ, nơi dữ liệu thường thưa và nhiễu hơn. "
        "Nói cách khác, hourly là cấp phù hợp để sửa target, còn daily là cấp phù hợp để ra quyết định forecast.",
    )

    doc.add_heading("3.4. Khôi phục Nhu cầu bằng Cửa sổ Mở rộng", level=2)
    add_paragraph(
        doc,
        "Trong train period, recovery được thực hiện bằng expanding window theo block tuần. Với mỗi block cần recover, mô hình chỉ được huấn luyện trên các quan sát non-stockout nằm trước block đó. "
        "Cơ chế này tránh việc dùng ngày tương lai để khôi phục nhu cầu của ngày quá khứ, đồng thời mô phỏng điều kiện triển khai thực tế: tại thời điểm hiện tại chỉ có dữ liệu quá khứ.",
    )
    add_figure(
        doc,
        "owner_expanding_window_recovery_detail.png",
        "Hình 3.2: Quy trình expanding-window recovery tránh leakage",
        "Hinh_3.2_Expanding_window_recovery_tranh_leakage.png",
    )
    add_table(
        doc,
        ["Block start", "Block end", "Training end", "Training rows", "Predicted rows"],
        df_rows("owner_recovery_blocks.csv", ["Block start", "Block end", "Training end", "Training rows", "Predicted rows"]),
        "Bảng 3.2: Các block recovery trong train period",
    )
    add_paragraph(
        doc,
        "Giai đoạn warm-up 14 ngày đầu được giữ observed để mô hình có đủ lịch sử cho lag, rolling và pattern tuần đầu tiên. "
        f"Sau warm-up, recovery được thực hiện qua {recovery.get('Recovery blocks', '')} block, mỗi block {recovery.get('Recovery block days', '')} ngày. "
        "Với validation và test, mô hình recovery được fit từ các quan sát non-stockout trong train period, sau đó áp dụng cho các giờ stockout ở phần sau.",
    )

    doc.add_heading("3.5. Calibration, Capping và Kiểm soát Chất lượng Imputation", level=2)
    add_paragraph(
        doc,
        "Imputation không được xem là ground truth tuyệt đối, mà là proxy cần được kiểm soát. Sau khi mô hình recovery dự đoán nhu cầu tiềm ẩn, phần lost demand chỉ được cộng thêm khi dự đoán lớn hơn observed sales. "
        "Phần cộng thêm này được calibration để giảm khuynh hướng phóng đại và được cap tại ngưỡng q90 của calibrated positive lost demand.",
    )
    add_table(
        doc,
        ["Chỉ tiêu", "Giá trị"],
        [
            ["Observed sales", fmt_float(recovery.get("Observed sales", ""), 2)],
            ["Recovered latent demand", fmt_float(recovery.get("Recovered latent demand", ""), 2)],
            ["Recovered lost demand", fmt_float(recovery.get("Recovered lost demand", ""), 2)],
            ["Recovered lift over observed", fmt_pct(recovery.get("Recovered lift over observed", ""))],
            ["Stockout row rate", fmt_pct(recovery.get("Stockout row rate", ""))],
            ["Calibration factor", fmt_float(recovery.get("Imputation calibration factor", ""), 4)],
            ["Calibration source", recovery.get("Calibration source", "")],
            ["Lost demand cap", fmt_float(recovery.get("Lost demand cap value", ""), 4)],
            ["Lost demand cap source", recovery.get("Lost demand cap source", "")],
        ],
        "Bảng 3.3: Tóm tắt kết quả và tham số recovery",
    )
    add_table(
        doc,
        ["Scenario", "Cap value", "Recovered demand", "Recovered lost demand", "Recovered lift over observed", "Share of original recovered lost demand"],
        df_rows(
            "imputation_cap_sensitivity.csv",
            ["Scenario", "Cap value", "Recovered demand", "Recovered lost demand", "Recovered lift over observed", "Share of original recovered lost demand"],
        ),
        "Bảng 3.4: Phân tích nhạy cảm theo ngưỡng cap lost demand",
    )
    add_figure(
        doc,
        "imputation_daily_uplift.png",
        "Hình 3.3: Mức uplift sau imputation theo ngày",
        "Hinh_3.3_Uplift_sau_imputation_theo_ngay.png",
    )
    add_paragraph(
        doc,
        "Việc dùng q90 là lựa chọn cân bằng: q50 có thể quá bảo thủ và bỏ sót nhiều nhu cầu bị che khuất, trong khi q100 giữ toàn bộ phần lost demand raw và dễ làm tổng demand tăng quá mức. "
        "Ở cấu hình q90, recovered demand tăng khoảng "
        f"{fmt_pct(recovery.get('Recovered lift over observed', ''))} so với observed sales, mức tăng đủ phản ánh stockout nhưng vẫn được kiểm soát bằng cap.",
    )

    doc.add_heading("3.6. Xác định Mục tiêu Dự báo Cấp độ Ngày", level=2)
    add_paragraph(
        doc,
        "Sau khi hoàn tất recovery ở cấp giờ, dữ liệu được tổng hợp lên cấp ngày theo từng series_id. Target chính của bài toán là tổng recovered demand trong 7 ngày tiếp theo, ký hiệu là target_next7_recovered_daily. "
        "Cách định nghĩa này phù hợp với bài toán vận hành vì người quản lý thường cần ước lượng tổng nhu cầu ngắn hạn để lên kế hoạch tồn kho, thay vì chỉ cần dự báo từng giờ riêng lẻ.",
    )
    add_paragraph(
        doc,
        "Về mặt công thức, tại forecast origin t của một chuỗi, target được tính bằng tổng nhu cầu đã recover từ ngày t+1 đến t+7. "
        "Việc dùng cửa sổ tương lai 7 ngày giúp bài toán gắn với quyết định replenishment theo tuần, đồng thời tận dụng seasonality tuần đã quan sát ở Chương 2.",
    )

    doc.add_heading("3.7. Lựa chọn Đặc trưng và Mô hình", level=2)
    add_paragraph(
        doc,
        "Feature engineering được xây dựng theo ba nhóm chính: đặc trưng định danh và lịch, đặc trưng lag/rolling của doanh số hoặc demand, và đặc trưng ngoại sinh như khuyến mãi, thời tiết, ngày lễ. "
        "Các biến lag và rolling phản ánh kết quả EDA về autocorrelation và seasonality; các biến lịch giúp mô hình học nhịp ngày/tuần; các biến ngoại sinh giúp bổ sung thông tin về điều kiện bán hàng.",
    )
    add_paragraph(
        doc,
        "Để xử lý rõ hơn hai điểm mù của observed sales, đồ án bổ sung hai nhóm đặc trưng giải thích được. "
        "Nhóm thứ nhất là tín hiệu thay thế trong cùng cửa hàng và cùng nhóm hàng, gồm tổng sales của các sản phẩm peer tại cùng thời điểm, stockout rate của nhóm peer và tốc độ tăng sales của nhóm peer. "
        "Nhóm thứ hai là tín hiệu đà bán trước stockout, gồm tỷ lệ rolling 3 giờ so với rolling 24 giờ, tỷ lệ sale giờ gần nhất so với nền 24 giờ và momentum giữa các cửa sổ ngắn. "
        "Nhờ vậy, recovery không chỉ dựa vào lịch sử riêng của một sản phẩm, mà còn quan sát được bối cảnh xung quanh: sản phẩm tương tự có đang bán mạnh lên không và bản thân sản phẩm có tăng tốc trước khi mất quan sát không.",
    )
    add_table(
        doc,
        ["Nhóm đặc trưng", "Ví dụ", "Vai trò"],
        [
            ["Định danh", "city_id, store_id, product_id, category_id", "Cho phép global model học khác biệt giữa cửa hàng và sản phẩm"],
            ["Lịch", "day_of_week, is_weekend, month, sin/cos calendar", "Biểu diễn mùa vụ và nhịp hoạt động"],
            ["Lag/Rolling", "lag 7/14/28, rolling mean/std/min/max", "Tóm tắt lịch sử nhu cầu gần và mùa vụ"],
            ["Stockout", "stockout rate, lag/rolling stockout", "Ghi nhận mức độ kiểm duyệt của observed sales"],
            ["Velocity/Momentum", "sale_velocity_ratio_3_24, sale_momentum_1_3", "Nhận diện nhu cầu tăng nhanh ngay trước khi xảy ra stockout"],
            ["Peer/Substitution", "peer_sales_same_group, peer_stockout_rate_same_group", "Bắt tín hiệu sản phẩm thay thế cùng cửa hàng và cùng nhóm hàng"],
            ["Ngoại sinh", "discount, holiday, activity, weather", "Bổ sung tín hiệu khuyến mãi, ngày lễ và thời tiết"],
        ],
        "Bảng 3.5: Các nhóm đặc trưng sử dụng trong mô hình",
    )
    add_paragraph(
        doc,
        "Mô hình chính được chọn là LightGBM dạng global model. Lý do là LightGBM xử lý tốt dữ liệu bảng có nhiều đặc trưng phi tuyến, chạy nhẹ hơn các mô hình deep learning phức tạp, đồng thời học được pattern chung từ nhiều chuỗi thay vì phải fit một mô hình riêng cho từng series. "
        "Trong bối cảnh dữ liệu thưa, nhiều chuỗi và có stockout, đây là lựa chọn cân bằng giữa hiệu suất, tốc độ huấn luyện và khả năng giải thích qua feature importance.",
    )

    doc.add_heading("3.8. Các Mô hình Cơ sở và Khung Kết hợp Hybrid", level=2)
    add_paragraph(
        doc,
        "Baseline được thiết lập để kiểm tra liệu mô hình học máy có thật sự vượt qua các quy luật đơn giản của chuỗi thời gian hay không. Với dữ liệu này, seasonal naive 7-day là baseline quan trọng nhất vì Chương 2 đã chỉ ra seasonality tuần rõ ràng. "
        "Naive, rolling mean và mô hình observed-sales được dùng như các đối chứng bổ sung.",
    )
    add_table(
        doc,
        ["Nhóm mô hình", "Mục đích"],
        [
            ["Naive / rolling mean", "Đo mức hiệu quả của các quy tắc đơn giản không cần huấn luyện"],
            ["Seasonal naive 7-day", "Kiểm tra sức mạnh của pattern tuần"],
            ["Observed-sales LightGBM", "Đo rủi ro khi dự báo trực tiếp observed sales"],
            ["Recovered-demand LightGBM", "Đánh giá lợi ích của target sau recovery"],
            ["Seasonal-ML hybrid", "Kết hợp pattern tuần mạnh với phần hiệu chỉnh học máy"],
        ],
        "Bảng 3.6: Vai trò của các baseline và mô hình so sánh",
    )
    add_paragraph(
        doc,
        "Hybrid framework kết hợp seasonal naive và LightGBM trên recovered demand. Trọng số được chọn trên validation set, không chọn bằng test set. "
        "Theo bảng lựa chọn trọng số, cấu hình tốt nhất dùng LightGBM weight "
        f"{fmt_float(hybrid_selection['LightGBM weight'], 2)} và seasonal naive weight "
        f"{fmt_float(hybrid_selection['Seasonal naive weight'], 2)}. "
        "Cách này giữ lại pattern tuần mạnh, đồng thời cho phép mô hình học máy hiệu chỉnh theo các biến định danh, lịch, stockout và ngoại sinh.",
    )
    add_table(
        doc,
        ["LightGBM weight", "Seasonal naive weight", "Validation WAPE"],
        df_rows("owner_hybrid_blend_selection.csv", ["LightGBM weight", "Seasonal naive weight", "Validation WAPE"], max_rows=5),
        "Bảng 3.7: Một số cấu hình trọng số hybrid tốt nhất trên validation set",
    )

    doc.add_heading("3.9. Các Tiêu chí Đánh giá Hiệu suất", level=2)
    add_paragraph(
        doc,
        "Các chỉ số đánh giá được chọn để phù hợp với dữ liệu nhiều giá trị 0 và phân phối lệch phải. WAPE là chỉ số chính vì nó đo tổng sai số tuyệt đối so với tổng nhu cầu thực tế, ổn định hơn MAPE khi có nhiều quan sát nhỏ hoặc bằng 0. "
        "MAE và RMSE được dùng bổ sung để lần lượt phản ánh sai số tuyệt đối trung bình và mức phạt lớn hơn cho các lỗi lớn.",
    )
    add_table(
        doc,
        ["Metric", "Công thức/diễn giải", "Vai trò trong đồ án"],
        [
            ["MAE", "mean(|y - y_hat|)", "Đo sai số tuyệt đối trung bình"],
            ["RMSE", "sqrt(mean((y - y_hat)^2))", "Nhấn mạnh các lỗi dự báo lớn"],
            ["WAPE", "sum(|y - y_hat|) / sum(|y|)", "Metric chính, phù hợp với dữ liệu nhiều zero"],
            ["sMAPE", "Sai số phần trăm đối xứng", "Đối chứng tương đối khi quy mô chuỗi khác nhau"],
            ["WPE", "sum(y_hat - y) / sum(y)", "Đo xu hướng over-forecast hoặc under-forecast"],
        ],
        "Bảng 3.8: Các tiêu chí đánh giá mô hình",
    )
    add_paragraph(
        doc,
        "Ngoài đánh giá forecast, đồ án còn kiểm tra chất lượng recovery bằng pseudo-stockout validation và sensitivity theo cap. "
        "Cách đánh giá này giúp tránh tình huống mô hình forecasting cải thiện chỉ vì target imputed bị phóng đại hoặc phụ thuộc quá mạnh vào bước khôi phục.",
    )

    doc.add_heading("Kết luận Chương 3", level=2)
    add_paragraph(
        doc,
        "Chương 3 xây dựng phương pháp theo logic nhất quán với EDA ở Chương 2. Vì observed sales bị ảnh hưởng bởi stockout, đồ án không dự báo trực tiếp doanh số quan sát như mục tiêu cuối cùng. "
        "Thay vào đó, pipeline khôi phục nhu cầu ở cấp giờ bằng expanding window để tránh leakage, kiểm soát imputation bằng calibration và cap, sau đó dự báo target daily next-7 trên recovered demand.",
    )
    add_paragraph(
        doc,
        "Việc chọn LightGBM global model và hybrid seasonal-ML xuất phát từ đặc tính dữ liệu: nhiều chuỗi, dữ liệu thưa, seasonality tuần mạnh và cần một mô hình đủ nhẹ để triển khai trong phạm vi đồ án. "
        "Các baseline và metric được thiết lập để đánh giá không chỉ độ chính xác, mà còn kiểm tra liệu phương pháp có thật sự cải thiện so với dự báo trực tiếp observed sales hay không.",
    )

    saved_path = save_doc(doc, OUT_PATH)
    print(saved_path)


if __name__ == "__main__":
    build_doc()
