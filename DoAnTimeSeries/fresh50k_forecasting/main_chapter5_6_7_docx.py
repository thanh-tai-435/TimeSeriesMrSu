from __future__ import annotations

import csv
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TABLE_DIR = BASE_DIR / "outputs" / "tables"
FIGURE_DIR = BASE_DIR / "outputs" / "figures"
OUT_DIR = BASE_DIR / "deliverables" / "reports"
OUT_PATH = OUT_DIR / "Chuong_5_6_7_Chan_doan_Business_Tong_ket.docx"
IMAGE_OUT_DIR = OUT_DIR / "Chuong_5_6_7_Hinh_anh_goc"


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
        doc.save(fallback)
        return fallback


def set_default_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12 if style_name == "Normal" else 14)
        style.font.bold = style_name.startswith("Heading")
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)

    for style_name, italic, bold in [("Caption", True, False), ("Table Caption", False, True)]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.font.italic = italic
        style.font.bold = bold
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    doc.add_paragraph(caption, style="Table Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, export_filename: str, width: float = 6.2) -> None:
    path = FIGURE_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[Thiếu hình: {filename}]")
        return
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, IMAGE_OUT_DIR / export_filename)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")


def read_metric_table(filename: str) -> dict[str, str]:
    result: dict[str, str] = {}
    with (TABLE_DIR / filename).open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            metric_key = "Metric" if "Metric" in row else "metric"
            value_key = "Value" if "Value" in row else "value"
            result[row[metric_key]] = row[value_key]
    return result


def fmt_float(value: str | float, digits: int = 4) -> str:
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)


def fmt_pct(value: str | float, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def fmt_int(value: str | float) -> str:
    try:
        return f"{int(float(value)):,}"
    except Exception:
        return str(value)


def format_cell(col: str, value) -> str:
    if isinstance(value, float):
        lowered = col.lower()
        if "p-value" in lowered:
            return f"{value:.4f}"
        if any(key in lowered for key in ["wape", "smape", "wpe", "coverage", "lift", "share", "ratio", "rate"]):
            return fmt_pct(value)
        return fmt_float(value, 4)
    return str(value)


def df_rows(filename: str, columns: list[str], max_rows: int | None = None) -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / filename)
    if max_rows is not None:
        df = df.head(max_rows)
    return [[format_cell(col, row[col]) for col in columns] for _, row in df.iterrows()]


def selected_diagnostic_rows() -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_diagnostics.csv")
    keep = [
        "Recovered seasonal naive 7-day",
        "Observed-sales forecasting",
        "Recovered-demand forecasting",
        "Recovered seasonal-ML hybrid",
    ]
    df = df[df["Model"].isin(keep)].copy()
    df["rank"] = df["Model"].map({name: idx for idx, name in enumerate(keep)})
    df = df.sort_values("rank")
    cols = ["Model", "RMSE", "MAE", "WAPE", "WPE", "Residual mean", "Ljung-Box p-value", "Jarque-Bera p-value", "N"]
    return [[format_cell(col, row[col]) for col in cols] for _, row in df.iterrows()]


def selected_nonoverlap_rows() -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_nonoverlap_diagnostics.csv")
    keep = [
        "Recovered seasonal naive 7-day",
        "Observed-sales forecasting",
        "Recovered-demand forecasting",
        "Recovered seasonal-ML hybrid",
    ]
    df = df[df["Model"].isin(keep)].copy()
    df["rank"] = df["Model"].map({name: idx for idx, name in enumerate(keep)})
    df = df.sort_values("rank")
    cols = ["Model", "WAPE", "WPE", "Residual mean", "Ljung-Box p-value", "N"]
    return [[format_cell(col, row[col]) for col in cols] for _, row in df.iterrows()]


def interval_summary_rows() -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_prediction_intervals.csv")
    keep = [
        "Recovered seasonal naive 7-day",
        "Observed-sales forecasting",
        "Recovered-demand forecasting",
        "Recovered seasonal-ML hybrid",
    ]
    rows = []
    for model_name in keep:
        group = df[df["Model"] == model_name]
        if group.empty:
            continue
        coverage_80 = ((group["y_true_recovered"] >= group["lower_80"]) & (group["y_true_recovered"] <= group["upper_80"])).mean()
        coverage_95 = ((group["y_true_recovered"] >= group["lower_95"]) & (group["y_true_recovered"] <= group["upper_95"])).mean()
        width_80 = (group["upper_80"] - group["lower_80"]).mean()
        width_95 = (group["upper_95"] - group["lower_95"]).mean()
        rows.append(
            [
                model_name,
                fmt_pct(coverage_80),
                fmt_pct(coverage_95),
                fmt_float(width_80, 4),
                fmt_float(width_95, 4),
                fmt_int(len(group)),
            ]
        )
    return rows


def best_forecast_metrics() -> dict[str, float]:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_forecasting_comparison.csv")
    df = df[df["Evaluation target"] == "Recovered latent demand proxy"].copy()
    result = {}
    for model in ["Observed-sales forecasting", "Recovered-demand forecasting", "Recovered seasonal-ML hybrid", "Recovered seasonal naive 7-day"]:
        row = df[df["Model"] == model].iloc[0]
        result[f"{model} WAPE"] = float(row["WAPE"])
        result[f"{model} WPE"] = float(row["WPE"])
    return result


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if IMAGE_OUT_DIR.exists():
        shutil.rmtree(IMAGE_OUT_DIR)
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)

    recovery = read_metric_table("owner_latent_recovery_summary.csv")
    quality = read_metric_table("data_quality_summary.csv")
    metrics = best_forecast_metrics()
    observed_wape = metrics["Observed-sales forecasting WAPE"]
    hybrid_wape = metrics["Recovered seasonal-ML hybrid WAPE"]
    rel_improvement = (observed_wape - hybrid_wape) / observed_wape

    doc = Document()
    set_default_font(doc)

    title = doc.add_heading("CHƯƠNG 5: CHẨN ĐOÁN MÔ HÌNH VÀ ĐÁNH GIÁ ĐỘ BẤT ĐỊNH", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        "Sau khi so sánh hiệu suất dự báo ở Chương 4, Chương 5 tập trung kiểm tra độ tin cậy của mô hình. Với dữ liệu bán lẻ nhiều chuỗi, nhiều giá trị 0 và chịu ảnh hưởng stockout, một mô hình có WAPE thấp vẫn cần được đánh giá thêm qua phần dư, độ ổn định trên các cửa sổ không chồng lấn và khoảng dự báo. Mục tiêu của chương này không phải chứng minh mô hình hoàn hảo, mà là chỉ ra mô hình còn sai ở đâu và mức độ bất định cần được đọc như thế nào trong bối cảnh vận hành.",
    )

    doc.add_heading("5.1. Chẩn đoán Phần dư (Residual Diagnostics)", level=2)
    add_paragraph(
        doc,
        "Phần dư được tính trên recovered latent demand proxy, tức là target đã xử lý stockout ở cấp giờ và tổng hợp lên daily next-7-day demand. Vì target forecast là tổng 7 ngày tiếp theo, các forecast origin liên tiếp có target window chồng lấn mạnh. Điều này khiến phần dư có thể còn tự tương quan ngay cả khi mô hình đã học được phần lớn seasonality. Do đó, Ljung-Box và residual ACF được dùng như công cụ chẩn đoán, không phải điều kiện để loại bỏ mô hình.",
    )
    add_table(
        doc,
        ["Model", "RMSE", "MAE", "WAPE", "WPE", "Residual mean", "Ljung-Box p-value", "Jarque-Bera p-value", "N"],
        selected_diagnostic_rows(),
        "Bảng 5.1: Chẩn đoán phần dư trên recovered latent demand proxy",
    )
    add_figure(
        doc,
        "owner_two_stage_residual_acf.png",
        "Hình 5.1: ACF phần dư của mô hình two-stage tốt nhất",
        "Hinh_5.1_ACF_phan_du_two_stage.png",
    )
    add_figure(
        doc,
        "owner_two_stage_residual_distribution.png",
        "Hình 5.2: Phân phối phần dư của mô hình two-stage tốt nhất",
        "Hinh_5.2_Phan_phoi_phan_du_two_stage.png",
    )
    add_paragraph(
        doc,
        "Kết quả cho thấy hybrid seasonal-ML đạt WAPE thấp nhất trong nhóm mô hình chính, khoảng "
        f"{fmt_pct(hybrid_wape)}, nhưng phần dư chưa thể xem là white noise hoàn toàn. Đây là kết quả hợp lý với bài toán này vì dữ liệu có nhiều sản phẩm, nhiều cửa hàng, nhiều zero và chịu ảnh hưởng bởi stockout. Thay vì che giấu điểm này, đồ án xem đây là bằng chứng rằng mô hình đã giảm sai số vận hành nhưng vẫn còn pattern chưa học hết.",
    )

    doc.add_heading("5.2. Kiểm định Sự Bất tương trùng (Non-overlap Diagnostics)", level=2)
    add_paragraph(
        doc,
        "Do target chính là tổng nhu cầu 7 ngày tiếp theo, các forecast origin liên tiếp tạo ra target window chồng lấn. Ví dụ, origin hôm nay và origin ngày mai cùng chứa phần lớn các ngày tương lai giống nhau. Nếu chỉ đánh giá trên toàn bộ origin liên tiếp, số lượng mẫu kiểm định có thể lớn nhưng không hoàn toàn độc lập. Vì vậy, đồ án bổ sung non-overlap diagnostics bằng cách chỉ lấy các forecast origin cách nhau 7 ngày để giảm mức chồng lấn giữa các target window.",
    )
    add_table(
        doc,
        ["Model", "WAPE", "WPE", "Residual mean", "Ljung-Box p-value", "N"],
        selected_nonoverlap_rows(),
        "Bảng 5.2: Diagnostics trên các target window không chồng lấn",
    )
    add_paragraph(
        doc,
        "Kết quả non-overlap vẫn giữ cùng kết luận chính: hybrid seasonal-ML tiếp tục nằm trong nhóm tốt nhất và không bị đảo chiều so với seasonal naive hoặc recovered-demand LightGBM. Điều này làm tăng độ tin cậy của kết quả ở Chương 4, vì kết luận không chỉ đến từ việc tạo nhiều forecast origin có target window chồng lấn.",
    )

    doc.add_heading("5.3. Phân tích Khoảng Dự báo (Prediction Intervals) và Định lượng Độ Cậy", level=2)
    add_paragraph(
        doc,
        "Trong ứng dụng replenishment, point forecast không đủ để ra quyết định nhập hàng. Nhà vận hành cần biết vùng dao động hợp lý của nhu cầu để cân bằng giữa thiếu hàng và tồn kho dư. Vì vậy, đồ án tạo khoảng dự báo dựa trên phân phối sai số validation. Cách làm này đơn giản hơn mô hình Bayesian hoặc quantile model đầy đủ, nhưng phù hợp với phạm vi đồ án và minh bạch khi giải thích.",
    )
    add_table(
        doc,
        ["Model", "80% coverage", "95% coverage", "Mean 80% width", "Mean 95% width", "N"],
        interval_summary_rows(),
        "Bảng 5.3: Độ phủ và độ rộng trung bình của khoảng dự báo",
    )
    add_figure(
        doc,
        "owner_two_stage_forecast_interval.png",
        "Hình 5.3: Forecast interval trên recovered latent demand",
        "Hinh_5.3_Khoang_du_bao_two_stage.png",
    )
    add_paragraph(
        doc,
        "Khoảng dự báo 95% có độ phủ cao hơn 80% nhưng rộng hơn đáng kể. Trong thực tế, mức khoảng sử dụng nên phụ thuộc vào chi phí hết hàng và chi phí tồn kho. Với nhóm hàng tươi sống, tồn kho dư có thể gây hao hụt, nhưng hết hàng cũng làm mất doanh số và làm sai lệch dữ liệu học trong tương lai. Do đó, khoảng dự báo nên được dùng như công cụ hỗ trợ quyết định, không phải cam kết chắc chắn về nhu cầu.",
    )

    title = doc.add_heading("CHƯƠNG 6: GÓC NHÌN KINH DOANH VÀ ỨNG DỤNG THỰC TẾ (BUSINESS INSIGHTS)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        "Các chương trước tập trung vào phương pháp và đánh giá mô hình. Chương 6 chuyển các kết quả đó thành góc nhìn kinh doanh: stockout làm sai lệch cách đọc nhu cầu như thế nào, lost sales có thể được ước lượng ra sao, seasonality tuần hỗ trợ vận hành như thế nào và mô hình có thể được đưa vào quy trình nhập hàng theo cách nào.",
    )

    doc.add_heading("6.1. Hiện tượng đứt hàng (Stockout) làm sai lệch số liệu nhu cầu như thế nào", level=2)
    add_paragraph(
        doc,
        "Khi một sản phẩm hết hàng, observed sales không còn là demand mà trở thành demand bị giới hạn bởi tồn kho. Nếu hệ thống forecast học trực tiếp từ observed sales, các ngày/giờ stockout có thể bị hiểu nhầm là nhu cầu thấp. Về dài hạn, vòng lặp này gây under-forecast: vì dự báo thấp nên nhập ít, vì nhập ít nên tiếp tục hết hàng, rồi dữ liệu mới lại càng củng cố nhận định sai rằng nhu cầu thấp.",
    )
    add_figure(
        doc,
        "stockout_rate_over_time.png",
        "Hình 6.1: Stockout rate theo thời gian",
        "Hinh_6.1_Stockout_rate_theo_thoi_gian.png",
    )
    add_paragraph(
        doc,
        "Trong dữ liệu mẫu, stockout row rate đạt khoảng "
        f"{fmt_pct(recovery.get('Stockout row rate', ''))}. Mức này đủ lớn để ảnh hưởng trực tiếp đến target dự báo, do đó xử lý stockout không phải bước phụ mà là phần trung tâm của bài toán.",
    )

    doc.add_heading("6.2. Tính toán doanh số bị mất (Lost Sales) từ dữ liệu nhu cầu ẩn", level=2)
    add_paragraph(
        doc,
        "Lost sales được ước lượng bằng phần chênh lệch dương giữa recovered demand và observed sales tại các giờ stockout, sau calibration và capping. Cách tính này không khẳng định biết chính xác từng giao dịch bị mất, nhưng cung cấp một proxy có kiểm soát để định lượng mức nhu cầu bị che khuất.",
    )
    add_table(
        doc,
        ["Chỉ tiêu", "Giá trị"],
        [
            ["Observed sales", fmt_float(recovery.get("Observed sales", ""), 2)],
            ["Recovered latent demand", fmt_float(recovery.get("Recovered latent demand", ""), 2)],
            ["Recovered lost demand", fmt_float(recovery.get("Recovered lost demand", ""), 2)],
            ["Recovered lift over observed", fmt_pct(recovery.get("Recovered lift over observed", ""))],
            ["Lost demand cap source", recovery.get("Lost demand cap source", "")],
        ],
        "Bảng 6.1: Ước lượng lost sales từ recovered latent demand",
    )
    add_figure(
        doc,
        "imputation_daily_uplift.png",
        "Hình 6.2: Uplift recovered demand theo ngày",
        "Hinh_6.2_Uplift_recovered_demand_theo_ngay.png",
    )
    add_paragraph(
        doc,
        "Kết quả cho thấy recovered latent demand cao hơn observed sales khoảng "
        f"{fmt_pct(recovery.get('Recovered lift over observed', ''))}. Đây là phần nhu cầu có thể bị bỏ sót nếu doanh nghiệp chỉ nhìn vào doanh số ghi nhận. Với hệ thống vận hành, con số này giúp ưu tiên kiểm tra các nhóm hàng/cửa hàng có lost sales cao thay vì chỉ tối ưu theo sales đã quan sát.",
    )

    doc.add_heading("6.3. Tối ưu hóa Vận hành dựa trên Tính Mùa vụ theo Tuần", level=2)
    add_paragraph(
        doc,
        "Kết quả EDA và mô hình đều cho thấy seasonality tuần là tín hiệu rất mạnh. Seasonal naive 7-day đạt WAPE khoảng "
        f"{fmt_pct(metrics['Recovered seasonal naive 7-day WAPE'])}, gần ngang với mô hình học máy. Điều này có ý nghĩa kinh doanh rõ ràng: lịch nhập hàng không nên chỉ phản ứng theo vài giờ hoặc vài ngày gần nhất, mà cần nhìn cùng kỳ tuần trước và các tuần trước đó.",
    )
    add_figure(
        doc,
        "spectrum_daily_recovered_demand.png",
        "Hình 6.3: Chu kỳ tuần trong recovered daily demand",
        "Hinh_6.3_Chu_ky_tuan_recovered_daily_demand.png",
    )
    add_paragraph(
        doc,
        "Trong thực tế, điều này gợi ý nên xây dựng các rule vận hành theo ngày trong tuần: nhu cầu thứ Hai không nhất thiết giống Chủ nhật, và nhu cầu cuối tuần có thể cần kế hoạch nhập hàng khác. Seasonal baseline cũng nên được giữ như benchmark vận hành lâu dài, vì một mô hình phức tạp nhưng không vượt seasonal naive thì chưa chứng minh được giá trị triển khai.",
    )

    doc.add_heading("6.4. Tối ưu vận hành dựa trên biến động mua sắm theo tuần", level=2)
    add_paragraph(
        doc,
        "Ngoài mùa vụ tuần ổn định, dữ liệu còn có biến động ngắn hạn trước và trong stockout. Các feature velocity/momentum giúp phát hiện trường hợp tốc độ bán gần đây cao hơn nền lịch sử, trong khi peer/substitution feature giúp kiểm tra bối cảnh sản phẩm cùng nhóm. Điều này quan trọng vì nhu cầu có thể không biến mất khi một sản phẩm hết hàng; nó có thể chuyển sang sản phẩm thay thế.",
    )
    add_table(
        doc,
        ["Segment", "Rows", "Velocity 3h/24h", "Momentum 3h-6h", "Peer sales", "Peer velocity 3h/24h", "Peer stockout rate"],
        df_rows(
            "stockout_substitution_velocity_diagnostics.csv",
            [
                "Segment",
                "Rows",
                "Mean sale velocity ratio 3h/24h",
                "Mean sale momentum 3h-6h",
                "Mean peer sales same group",
                "Mean peer velocity ratio 3h/24h",
                "Mean peer stockout rate",
            ],
        ),
        "Bảng 6.2: Tín hiệu velocity và peer/substitution quanh stockout",
    )
    add_figure(
        doc,
        "stockout_substitution_velocity_diagnostics.png",
        "Hình 6.4: So sánh velocity và peer signal giữa stockout và non-stockout",
        "Hinh_6.4_Velocity_peer_signal_stockout.png",
    )
    add_paragraph(
        doc,
        "Kết quả không nên được diễn giải rằng mọi stockout đều do demand surge. Ý nghĩa chính là hệ thống forecast nên có khả năng quan sát các tín hiệu trước stockout và quanh sản phẩm thay thế. Đây là cách biến dữ liệu thành cảnh báo vận hành: khi một sản phẩm bán nhanh hơn nền gần đây hoặc nhóm peer biến động mạnh, cửa hàng nên kiểm tra tồn kho sớm hơn.",
    )

    doc.add_heading("6.5. Ứng dụng mô hình dự báo vào quy trình nhập hàng (Replenishment)", level=2)
    add_paragraph(
        doc,
        "Mô hình cuối cùng dự báo tổng recovered demand trong 7 ngày tiếp theo ở cấp series_id. Trong quy trình replenishment, forecast này có thể được dùng như baseline nhu cầu tuần tới. Một quy trình triển khai tối giản gồm: cập nhật dữ liệu bán hàng và stockout theo giờ; chạy recovery để ước lượng demand bị che; tổng hợp lên daily; tạo forecast next-7-day; sau đó chuyển forecast thành đề xuất nhập hàng có xét tồn kho hiện tại, lead time, minimum order quantity và safety stock.",
    )
    add_table(
        doc,
        ["Bước", "Đầu vào", "Đầu ra vận hành"],
        [
            ["1. Cập nhật dữ liệu", "Observed sales, stock status, promotion, calendar", "Bảng hourly mới nhất"],
            ["2. Recovery", "Hourly features và stockout flag", "Recovered demand proxy"],
            ["3. Forecast", "Daily recovered demand và lag/rolling features", "Next-7-day demand forecast"],
            ["4. Quy đổi nhập hàng", "Forecast, tồn kho, lead time, safety stock", "Đề xuất lượng nhập"],
            ["5. Theo dõi", "Actual sales, stockout, forecast error", "Cảnh báo bias và drift"],
        ],
        "Bảng 6.3: Quy trình ứng dụng forecast vào replenishment",
    )
    add_figure(
        doc,
        "owner_two_stage_forecast_comparison.png",
        "Hình 6.5: Forecast next-7-day demand theo mô hình two-stage",
        "Hinh_6.5_Forecast_next7_demand_two_stage.png",
    )
    add_paragraph(
        doc,
        "Trong triển khai thực tế, doanh nghiệp không nên dùng point forecast một cách cứng nhắc. Với nhóm hàng có chi phí stockout cao, có thể chọn mức nhập gần upper bound của khoảng dự báo; với nhóm hàng tươi sống dễ hư hỏng, có thể dùng mức gần median/point forecast và tăng tần suất cập nhật.",
    )

    doc.add_heading("6.6. Khuyến nghị các bước triển khai mô hình vào thực tế", level=2)
    add_paragraph(
        doc,
        "Để triển khai ngoài môi trường đồ án, mô hình nên được đưa vào hệ thống theo từng bước thay vì áp dụng ngay toàn bộ. Bước đầu tiên là chạy song song với quy trình hiện tại để đo forecast error, lost sales proxy và stockout rate theo cửa hàng/nhóm hàng. Bước thứ hai là thử nghiệm A/B trên một số nhóm hàng có stockout cao. Bước thứ ba là theo dõi drift và hiệu chỉnh lại mô hình định kỳ khi có thay đổi về vận hành, khuyến mãi hoặc hành vi mua hàng.",
    )
    add_table(
        doc,
        ["Khuyến nghị", "Ý nghĩa"],
        [
            ["Theo dõi WAPE và WPE theo store/category", "Phát hiện nhóm bị under-forecast hoặc over-forecast"],
            ["Theo dõi stockout rate sau triển khai", "Kiểm tra forecast có giúp giảm đứt hàng không"],
            ["Giữ seasonal naive làm benchmark thường trực", "Tránh triển khai model phức tạp nhưng không vượt baseline"],
            ["Kiểm soát imputation bằng cap/sensitivity", "Tránh để recovered demand bị phóng đại"],
            ["Tái huấn luyện định kỳ", "Thích ứng với thay đổi seasonality, khuyến mãi và danh mục hàng"],
        ],
        "Bảng 6.4: Khuyến nghị triển khai mô hình vào thực tế",
    )

    title = doc.add_heading("CHƯƠNG 7: TỔNG KẾT VÀ HƯỚNG PHÁT TRIỂN", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("7.1. Kết luận Chung", level=2)
    add_paragraph(
        doc,
        "Đồ án đã xây dựng một pipeline dự báo nhu cầu cho dữ liệu bán lẻ có stockout theo hướng two-stage. Thay vì forecast trực tiếp observed sales, pipeline khôi phục latent demand ở cấp giờ, kiểm soát leakage bằng expanding window, hiệu chỉnh imputation bằng calibration và cap, sau đó tổng hợp lên daily demand để dự báo tổng nhu cầu 7 ngày tiếp theo.",
    )
    add_paragraph(
        doc,
        "Kết quả chính cho thấy recovered latent demand cao hơn observed sales khoảng "
        f"{fmt_pct(recovery.get('Recovered lift over observed', ''))}, phản ánh phần nhu cầu có thể bị che bởi stockout. Ở bước forecast, hybrid seasonal-ML đạt WAPE khoảng "
        f"{fmt_pct(hybrid_wape)}, tốt hơn observed-sales forecasting với WAPE khoảng {fmt_pct(observed_wape)}. Mức cải thiện tương đối khoảng {fmt_pct(rel_improvement)} cho thấy xử lý stockout trước khi forecast có giá trị thực nghiệm rõ ràng.",
    )
    add_paragraph(
        doc,
        "Một kết luận quan trọng khác là seasonal naive 7-day rất mạnh. Điều này không làm giảm giá trị của mô hình học máy, mà giúp framing đúng hơn: mô hình cuối cùng nên giữ pattern tuần làm nền và dùng LightGBM để hiệu chỉnh theo lịch sử demand, stockout, định danh sản phẩm-cửa hàng và các biến bổ sung.",
    )

    doc.add_heading("7.2. Các Hạn chế Tồn đọng của Đề tài", level=2)
    add_paragraph(
        doc,
        "Hạn chế đầu tiên là recovered demand vẫn là proxy, không phải ground truth tuyệt đối. Dù đã có pseudo-stockout validation, calibration và sensitivity, không thể biết chính xác từng khách hàng đã muốn mua bao nhiêu khi sản phẩm hết hàng. Vì vậy, mọi kết luận về lost sales cần được hiểu như ước lượng có kiểm soát.",
    )
    add_paragraph(
        doc,
        "Hạn chế thứ hai là residual diagnostics cho thấy phần dư chưa hoàn toàn độc lập và chuẩn. Đây là vấn đề thường gặp với dữ liệu bán lẻ nhiều chuỗi, nhiều zero và target 7 ngày chồng lấn. Mô hình đã giảm sai số nhưng vẫn còn pattern chưa học hết.",
    )
    add_paragraph(
        doc,
        "Hạn chế thứ ba là peer/substitution feature đã được bổ sung nhưng chưa trở thành tín hiệu mạnh nhất trong sample hiện tại. Điều này có thể do dữ liệu sample 10%, cách định nghĩa nhóm thay thế còn đơn giản hoặc hiệu ứng thay thế chỉ rõ ở một số danh mục cụ thể. Do đó, không nên overclaim rằng mô hình đã giải quyết hoàn toàn bài toán substitution.",
    )
    add_paragraph(
        doc,
        "Hạn chế cuối cùng là đồ án ưu tiên mô hình nhẹ, dễ giải thích và tái lập. Các mô hình deep imputation như SAITS, CSDI, TimesNet hoặc ImputeFormer chưa được triển khai vì yêu cầu tuning và tài nguyên lớn hơn phạm vi đồ án.",
    )

    doc.add_heading("7.3. Đề xuất Hướng Nghiên cứu và Phát triển Tương lai", level=2)
    add_paragraph(
        doc,
        "Hướng phát triển đầu tiên là cải thiện latent demand recovery bằng các mô hình chuyên cho missing/censored time series, ví dụ SAITS, CSDI, TimesNet hoặc ImputeFormer. Tuy nhiên, các mô hình này cần được so sánh nghiêm túc với baseline hiện tại bằng cùng split thời gian và cùng pseudo-stockout validation để tránh cải thiện giả do leakage hoặc over-imputation.",
    )
    add_paragraph(
        doc,
        "Hướng thứ hai là mô hình hóa substitution sâu hơn. Thay vì chỉ dùng peer sales cùng store/category, có thể xây dựng graph sản phẩm dựa trên độ tương quan nhu cầu, cùng giỏ hàng hoặc quan hệ thay thế/bổ sung. Khi đó, stockout của một sản phẩm có thể được truyền tín hiệu sang các sản phẩm liên quan một cách có cấu trúc hơn.",
    )
    add_paragraph(
        doc,
        "Hướng thứ ba là mở rộng forecast từ point forecast sang probabilistic forecasting. Với replenishment, phân phối nhu cầu và quantile forecast có giá trị trực tiếp hơn một con số trung bình, vì quyết định nhập hàng phụ thuộc vào trade-off giữa thiếu hàng và tồn kho dư.",
    )
    add_paragraph(
        doc,
        "Hướng cuối cùng là triển khai đánh giá thực nghiệm theo business metric: giảm stockout rate, giảm lost sales proxy, giảm hàng hủy do tồn kho dư và cải thiện service level. Khi đó, mô hình không chỉ được đánh giá bằng WAPE mà còn bằng tác động thực tế đến vận hành.",
    )

    saved_path = save_doc(doc, OUT_PATH)
    print(saved_path)


if __name__ == "__main__":
    build_doc()
