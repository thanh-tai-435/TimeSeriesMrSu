from __future__ import annotations

import csv
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TABLE_DIR = BASE_DIR / "outputs" / "tables"
FIGURE_DIR = BASE_DIR / "outputs" / "figures"
OUT_DIR = BASE_DIR / "deliverables" / "reports"
OUT_PATH = OUT_DIR / "Chuong_6_Business_Insights_Ung_dung_thuc_te.docx"
IMAGE_OUT_DIR = OUT_DIR / "Chuong_6_Hinh_anh_goc"


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
        doc.save(fallback)
        return fallback


def set_default_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12 if style_name == "Normal" else 14)
        style.font.bold = style_name.startswith("Heading")
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)

    for style_name, italic, bold in [("Caption", True, False), ("Table Caption", False, True)]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.font.italic = italic
        style.font.bold = bold
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    doc.add_paragraph(caption, style="Table Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, export_filename: str, width: float = 6.2) -> None:
    path = FIGURE_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[Thiếu hình: {filename}]")
        return
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, IMAGE_OUT_DIR / export_filename)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")


def read_metric_table(filename: str) -> dict[str, str]:
    result: dict[str, str] = {}
    with (TABLE_DIR / filename).open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            metric_key = "Metric" if "Metric" in row else "metric"
            value_key = "Value" if "Value" in row else "value"
            result[row[metric_key]] = row[value_key]
    return result


def fmt_float(value: str | float, digits: int = 4) -> str:
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)


def fmt_pct(value: str | float, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def fmt_int(value: str | float) -> str:
    try:
        return f"{int(float(value)):,}"
    except Exception:
        return str(value)


def format_cell(col: str, value) -> str:
    if isinstance(value, float):
        lowered = col.lower()
        if any(key in lowered for key in ["wape", "smape", "wpe", "lift", "share", "ratio", "rate"]):
            return fmt_pct(value)
        return fmt_float(value, 4)
    return str(value)


def table_rows(filename: str, columns: list[str], max_rows: int | None = None) -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / filename)
    if max_rows is not None:
        df = df.head(max_rows)
    return [[format_cell(col, row[col]) for col in columns] for _, row in df.iterrows()]


def forecast_metrics() -> dict[str, float]:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_forecasting_comparison.csv")
    df = df[df["Evaluation target"] == "Recovered latent demand proxy"].copy()
    result = {}
    for model in ["Observed-sales forecasting", "Recovered-demand forecasting", "Recovered seasonal-ML hybrid", "Recovered seasonal naive 7-day"]:
        row = df[df["Model"] == model].iloc[0]
        result[f"{model} WAPE"] = float(row["WAPE"])
        result[f"{model} WPE"] = float(row["WPE"])
    return result


def top_feature_rows(max_rows: int = 12) -> list[list[str]]:
    df = pd.read_csv(TABLE_DIR / "owner_recovereddemand_forecasting_feature_importance.csv").head(max_rows)
    return [[str(row["feature"]), fmt_int(row["importance"])] for _, row in df.iterrows()]


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if IMAGE_OUT_DIR.exists():
        shutil.rmtree(IMAGE_OUT_DIR)
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)

    recovery = read_metric_table("owner_latent_recovery_summary.csv")
    metrics = forecast_metrics()
    observed_wape = metrics["Observed-sales forecasting WAPE"]
    hybrid_wape = metrics["Recovered seasonal-ML hybrid WAPE"]
    seasonal_wape = metrics["Recovered seasonal naive 7-day WAPE"]
    rel_gain_vs_observed = (observed_wape - hybrid_wape) / observed_wape

    doc = Document()
    set_default_font(doc)

    title = doc.add_heading("CHƯƠNG 6: GÓC NHÌN KINH DOANH VÀ ỨNG DỤNG THỰC TẾ (BUSINESS INSIGHTS)", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        "Các chương trước đã chứng minh về mặt kỹ thuật rằng observed sales trong dữ liệu bán lẻ có thể bị kiểm duyệt bởi stockout, và pipeline two-stage giúp giảm sai số khi dự báo recovered demand. Chương 6 chuyển kết quả đó thành góc nhìn kinh doanh. Trọng tâm của chương không phải là giới thiệu thêm mô hình mới, mà là trả lời câu hỏi: các kết quả forecast và recovery có thể giúp nhà bán lẻ ra quyết định tốt hơn ở đâu.",
    )
    add_paragraph(
        doc,
        "Trong bối cảnh hàng tươi sống, dự báo nhu cầu không chỉ phục vụ mục tiêu tăng doanh số. Một forecast tốt còn giúp cân bằng giữa hai rủi ro trái chiều: nhập thiếu gây stockout và mất doanh số, nhập dư gây tồn kho, giảm chất lượng hoặc hủy hàng. Vì vậy, business insight của đồ án được xây dựng quanh ba lớp: nhận diện sai lệch do stockout, định lượng phần lost sales có thể bị che khuất, và dùng seasonality cùng forecast interval để hỗ trợ replenishment.",
    )
    add_paragraph(
        doc,
        "Một điểm cần nhấn mạnh là recovered demand trong đồ án không được xem là sự thật tuyệt đối. Đây là một proxy có kiểm soát bằng expanding window, calibration, cap và sensitivity. Do đó, các khuyến nghị kinh doanh trong chương này nên được hiểu là hướng ra quyết định dựa trên bằng chứng định lượng, không phải một quy tắc tự động thay thế hoàn toàn kinh nghiệm vận hành.",
    )

    doc.add_heading("6.1. Hiện tượng đứt hàng (Stockout) làm sai lệch số liệu nhu cầu như thế nào", level=2)
    add_paragraph(
        doc,
        "Trong bán lẻ, observed sales chỉ đo được lượng hàng đã bán, không đo được toàn bộ lượng khách hàng muốn mua. Khi sản phẩm còn hàng đầy đủ, observed sales có thể gần với demand. Nhưng khi xảy ra stockout, observed sales bị chặn bởi tồn kho. Lúc này, doanh số thấp không nhất thiết phản ánh nhu cầu thấp; nó có thể phản ánh việc cửa hàng không còn hàng để bán.",
    )
    add_paragraph(
        doc,
        "Nếu hệ thống forecast học trực tiếp trên observed sales, stockout sẽ tạo ra một vòng lặp sai lệch. Giai đoạn hết hàng làm doanh số ghi nhận thấp; mô hình học rằng nhu cầu thấp; forecast kỳ sau thấp; cửa hàng nhập ít; stockout tiếp tục xảy ra. Vòng lặp này đặc biệt nguy hiểm với hàng tươi sống vì quyết định nhập hàng thường ngắn hạn và nhạy với dữ liệu gần đây.",
    )
    add_table(
        doc,
        ["Chỉ tiêu", "Giá trị"],
        [
            ["Observed sales", fmt_float(recovery.get("Observed sales", ""), 2)],
            ["Recovered latent demand", fmt_float(recovery.get("Recovered latent demand", ""), 2)],
            ["Recovered lost demand", fmt_float(recovery.get("Recovered lost demand", ""), 2)],
            ["Recovered lift over observed", fmt_pct(recovery.get("Recovered lift over observed", ""))],
            ["Stockout row rate", fmt_pct(recovery.get("Stockout row rate", ""))],
        ],
        "Bảng 6.1: Quy mô ảnh hưởng của stockout lên nhu cầu quan sát",
    )
    add_figure(
        doc,
        "stockout_rate_over_time.png",
        "Hình 6.1: Tỷ lệ stockout theo thời gian",
        "Hinh_6.1_Ty_le_stockout_theo_thoi_gian.png",
    )
    add_paragraph(
        doc,
        "Với stockout row rate khoảng "
        f"{fmt_pct(recovery.get('Stockout row rate', ''))}, stockout không phải là hiện tượng hiếm. Nó đủ lớn để làm sai lệch target nếu không được xử lý. Đây là lý do business framing của đồ án không phải forecast sales đơn thuần, mà là forecast demand sau khi giảm bias do đứt hàng.",
    )
    add_paragraph(
        doc,
        "Từ góc nhìn quản trị, điều này thay đổi cách đọc báo cáo doanh số. Một cửa hàng có doanh số thấp ở một sản phẩm chưa chắc là cửa hàng có nhu cầu thấp; có thể đó là cửa hàng thường xuyên thiếu hàng. Vì vậy, dashboard vận hành nên hiển thị observed sales cùng stockout rate và recovered lost demand, thay vì chỉ xếp hạng cửa hàng/sản phẩm theo doanh số ghi nhận.",
    )

    doc.add_heading("6.2. Tính toán doanh số bị mất (Lost Sales) từ dữ liệu nhu cầu ẩn", level=2)
    add_paragraph(
        doc,
        "Lost sales trong đồ án được tính từ phần chênh lệch dương giữa recovered demand và observed sales tại các giờ stockout. Về mặt công thức, khi sản phẩm stockout, nếu mô hình recovery ước lượng nhu cầu cao hơn observed sales, phần chênh lệch đó được xem là lost demand. Sau đó lost demand được calibration và cap để tránh phóng đại. Với non-stockout rows, recovered demand giữ bằng observed sales.",
    )
    add_paragraph(
        doc,
        "Cách tính này có hai giá trị kinh doanh. Thứ nhất, nó biến stockout từ một cờ trạng thái thành một ước lượng định lượng về mức độ mất cơ hội bán hàng. Thứ hai, nó cho phép ưu tiên hành động theo impact: không phải mọi stockout đều quan trọng như nhau; stockout ở sản phẩm/cửa hàng có lost sales cao cần được ưu tiên hơn stockout ở chuỗi có nhu cầu thấp.",
    )
    add_table(
        doc,
        ["Scenario", "Cap value", "Recovered demand", "Recovered lost demand", "Recovered lift over observed", "Share of original recovered lost demand"],
        table_rows(
            "imputation_cap_sensitivity.csv",
            ["Scenario", "Cap value", "Recovered demand", "Recovered lost demand", "Recovered lift over observed", "Share of original recovered lost demand"],
        ),
        "Bảng 6.2: Độ nhạy của ước lượng lost sales theo ngưỡng cap",
    )
    add_figure(
        doc,
        "imputation_daily_uplift.png",
        "Hình 6.2: Uplift recovered demand theo ngày",
        "Hinh_6.2_Uplift_recovered_demand_theo_ngay.png",
    )
    add_paragraph(
        doc,
        "Bảng 6.2 cho thấy lựa chọn cap ảnh hưởng trực tiếp đến tổng lost sales. Nếu cap quá thấp như q50, recovery trở nên bảo thủ và có thể bỏ sót nhu cầu. Nếu cap q100, mô hình giữ toàn bộ phần lost demand raw nhưng rủi ro over-imputation cao hơn. Việc chọn q90 là một thỏa hiệp: đủ phản ánh stockout nhưng vẫn cắt bớt các dự đoán quá cực đoan.",
    )
    add_table(
        doc,
        ["series_id", "observed_sales", "recovered_demand", "recovered_lost_demand", "stockout_rate", "Recovered lift over observed"],
        table_rows(
            "imputation_series_uplift.csv",
            ["series_id", "observed_sales", "recovered_demand", "recovered_lost_demand", "stockout_rate", "Recovered lift over observed"],
            max_rows=10,
        ),
        "Bảng 6.3: Các chuỗi có uplift cao sau recovery",
    )
    add_figure(
        doc,
        "imputation_top_series_lift.png",
        "Hình 6.3: Các chuỗi có uplift cao sau imputation",
        "Hinh_6.3_Cac_chuoi_uplift_cao.png",
    )
    add_paragraph(
        doc,
        "Bảng 6.3 và Hình 6.3 giúp chuyển kết quả mô hình thành danh sách ưu tiên vận hành. Các chuỗi có recovered lift cao là những nơi observed sales có khả năng đang đánh giá thấp nhu cầu thật. Trong thực tế, nhóm này nên được kiểm tra về quy trình đặt hàng, tần suất bổ sung, tồn kho an toàn, lead time và độ chính xác của ghi nhận stock status.",
    )

    doc.add_heading("6.3. Tối ưu hóa Vận hành dựa trên Tính Mùa vụ theo Tuần", level=2)
    add_paragraph(
        doc,
        "Một insight nổi bật của đồ án là seasonality tuần rất mạnh. Seasonal naive 7-day đạt WAPE khoảng "
        f"{fmt_pct(seasonal_wape)}, gần sát mô hình hybrid. Điều này có nghĩa là lịch mua hàng có tính lặp lại cao theo tuần. Với vận hành bán lẻ, đây là tín hiệu quan trọng vì kế hoạch nhập hàng thường cũng được tổ chức theo nhịp ngày và tuần.",
    )
    add_figure(
        doc,
        "spectrum_daily_recovered_demand.png",
        "Hình 6.4: Chu kỳ tuần trong recovered daily demand",
        "Hinh_6.4_Chu_ky_tuan_recovered_daily_demand.png",
    )
    add_table(
        doc,
        ["Rank", "Period", "Power", "Period unit"],
        table_rows("spectrum_daily_recovered_top_peaks.csv", ["Rank", "Period", "Power", "Period unit"], max_rows=5),
        "Bảng 6.4: Các chu kỳ nổi bật trong recovered daily demand",
    )
    add_paragraph(
        doc,
        "Kết quả phổ tần số cho thấy chu kỳ gần 7 ngày là thành phần nổi bật. Vì vậy, trong quy trình vận hành, cùng kỳ tuần trước nên là điểm tham chiếu bắt buộc khi lập kế hoạch nhập hàng. Ví dụ, nhu cầu thứ Hai nên được so sánh với các thứ Hai trước đó hơn là chỉ so với Chủ nhật liền trước. Điều này đặc biệt quan trọng với hàng tươi sống, nơi hành vi mua sắm thường đi theo lịch sinh hoạt của khách hàng.",
    )
    add_paragraph(
        doc,
        "Seasonality tuần cũng giúp giải thích vì sao mô hình hybrid hoạt động tốt. Seasonal naive giữ lại nhịp tuần ổn định, còn LightGBM hiệu chỉnh theo lịch sử demand, stockout rate, store/product identity và các biến calendar khác. Về mặt business, đây là một cách kết hợp hợp lý giữa quy tắc vận hành dễ hiểu và mô hình học máy linh hoạt.",
    )

    doc.add_heading("6.4. Tối ưu vận hành dựa trên biến động mua sắm theo tuần", level=2)
    add_paragraph(
        doc,
        "Bên cạnh seasonality tuần, nhu cầu còn có biến động ngắn hạn. Một sản phẩm có thể bán nhanh bất thường trong vài giờ trước khi hết hàng, hoặc demand có thể dịch chuyển sang sản phẩm cùng nhóm khi sản phẩm chính stockout. Nếu mô hình chỉ nhìn lịch sử dài hạn của chính sản phẩm, hai tín hiệu này dễ bị bỏ qua.",
    )
    add_table(
        doc,
        ["Segment", "Rows", "Velocity 3h/24h", "Momentum 3h-6h", "Peer sales", "Peer velocity 3h/24h", "Peer stockout rate"],
        table_rows(
            "stockout_substitution_velocity_diagnostics.csv",
            [
                "Segment",
                "Rows",
                "Mean sale velocity ratio 3h/24h",
                "Mean sale momentum 3h-6h",
                "Mean peer sales same group",
                "Mean peer velocity ratio 3h/24h",
                "Mean peer stockout rate",
            ],
        ),
        "Bảng 6.5: Tín hiệu velocity và peer/substitution quanh stockout",
    )
    add_figure(
        doc,
        "stockout_substitution_velocity_diagnostics.png",
        "Hình 6.5: So sánh velocity và peer signal giữa stockout và non-stockout",
        "Hinh_6.5_Velocity_peer_signal_stockout.png",
    )
    add_paragraph(
        doc,
        "Bảng 6.5 không được diễn giải theo hướng mọi stockout đều do demand surge. Trái lại, kết quả cho thấy cần kiểm tra có hệ thống cả hai khả năng: một số stockout có thể xảy ra sau giai đoạn bán nhanh, trong khi một số stockout phản ánh vấn đề tồn kho kéo dài làm observed sales thấp. Việc bổ sung velocity/momentum giúp mô hình nhìn thấy đà bán gần nhất; việc bổ sung peer/substitution giúp mô hình có thêm bối cảnh về nhóm sản phẩm thay thế.",
    )
    add_paragraph(
        doc,
        "Về vận hành, các tín hiệu này có thể chuyển thành cảnh báo sớm. Nếu một sản phẩm bán nhanh hơn nền 24 giờ và tồn kho đang thấp, hệ thống có thể cảnh báo bổ sung hàng trước khi xảy ra stockout. Nếu một sản phẩm hết hàng và peer sales tăng bất thường, cửa hàng cần xem xét liệu demand đang chuyển sang sản phẩm thay thế hay không. Đây là insight quan trọng vì stockout không chỉ làm mất doanh số của một sản phẩm, mà còn làm méo cấu trúc nhu cầu của cả nhóm hàng.",
    )

    doc.add_heading("6.5. Ứng dụng mô hình dự báo vào quy trình nhập hàng (Replenishment)", level=2)
    add_paragraph(
        doc,
        "Mô hình cuối cùng dự báo tổng recovered demand trong 7 ngày tiếp theo ở cấp series_id. Đây là horizon phù hợp với replenishment theo tuần: đủ dài để hỗ trợ kế hoạch nhập hàng, nhưng không quá dài đến mức seasonality và biến động ngắn hạn bị làm mờ. Forecast không nên được dùng như con số nhập hàng trực tiếp; nó là đầu vào cho công thức replenishment có xét tồn kho hiện tại, lead time, mức trưng bày tối thiểu, lô đặt hàng tối thiểu và safety stock.",
    )
    add_table(
        doc,
        ["Bước", "Dữ liệu đầu vào", "Kết quả vận hành"],
        [
            ["1. Cập nhật dữ liệu", "Observed sales, stock status, promotion, weather, calendar", "Bảng hourly mới nhất"],
            ["2. Recovery", "Hourly features, stockout flag, peer/velocity signals", "Recovered demand proxy"],
            ["3. Aggregate", "Hourly recovered demand", "Daily recovered demand theo series_id"],
            ["4. Forecast", "Lag/rolling features và seasonal baseline", "Next-7-day demand forecast"],
            ["5. Quy đổi nhập hàng", "Forecast, tồn kho, lead time, safety stock", "Đề xuất lượng nhập"],
            ["6. Theo dõi sau triển khai", "Actual sales, stockout, forecast error", "Cảnh báo drift và bias"],
        ],
        "Bảng 6.6: Quy trình ứng dụng mô hình forecast vào replenishment",
    )
    add_figure(
        doc,
        "owner_two_stage_forecast_comparison.png",
        "Hình 6.6: So sánh forecast next-7-day demand giữa các mô hình",
        "Hinh_6.6_So_sanh_forecast_next7_demand.png",
    )
    add_table(
        doc,
        ["Feature", "Importance"],
        top_feature_rows(12),
        "Bảng 6.7: Các đặc trưng quan trọng trong recovered-demand forecasting",
    )
    add_figure(
        doc,
        "owner_recovereddemand_forecasting_feature_importance_top20.png",
        "Hình 6.7: Feature importance của recovered-demand LightGBM",
        "Hinh_6.7_Feature_importance_recovered_demand.png",
    )
    add_paragraph(
        doc,
        "Feature importance cho thấy stockout_rate và các lag/rolling của recovered demand là tín hiệu quan trọng. Điều này có ý nghĩa vận hành rõ ràng: lịch sử nhu cầu đã xử lý stockout và trạng thái stockout gần đây cần được đưa vào hệ thống replenishment. Nếu chỉ dùng observed sales, hệ thống có thể tiếp tục đánh giá thấp các chuỗi hay hết hàng.",
    )
    add_paragraph(
        doc,
        "Kết quả forecast cũng cho thấy hybrid seasonal-ML đạt WAPE khoảng "
        f"{fmt_pct(hybrid_wape)}, cải thiện khoảng {fmt_pct(rel_gain_vs_observed)} so với observed-sales forecasting. Trong thực tế, mức cải thiện này nên được quy đổi sang tác động kinh doanh như giảm stockout rate, giảm lost sales proxy hoặc cải thiện service level, thay vì chỉ dừng ở chỉ số mô hình.",
    )

    doc.add_heading("6.6. Khuyến nghị các bước triển khai mô hình vào thực tế", level=2)
    add_paragraph(
        doc,
        "Để triển khai thực tế, không nên đưa mô hình vào quyết định nhập hàng một cách đột ngột. Cách an toàn hơn là triển khai theo từng lớp. Giai đoạn đầu, mô hình chạy song song với quy trình hiện tại để tạo forecast, lost sales proxy và cảnh báo stockout nhưng chưa tự động thay đổi lượng nhập. Giai đoạn sau, chọn một nhóm cửa hàng/sản phẩm có stockout cao để thử nghiệm có kiểm soát. Chỉ khi kết quả vận hành cải thiện ổn định mới mở rộng phạm vi.",
    )
    add_table(
        doc,
        ["Nhóm chỉ số", "Cách theo dõi", "Mục đích"],
        [
            ["Forecast accuracy", "WAPE, MAE, WPE theo store/category", "Kiểm tra mô hình có bias ở nhóm cụ thể không"],
            ["Stockout impact", "Stockout rate, recovered lost demand", "Đo mức giảm đứt hàng và lost sales"],
            ["Inventory risk", "Tồn kho cuối ngày, hàng hủy, ngày tồn kho", "Tránh nhập dư khi tối ưu chống stockout"],
            ["Service level", "Tỷ lệ ngày/giờ còn hàng", "Đánh giá trải nghiệm khách hàng"],
            ["Model drift", "Sai số theo tuần, thay đổi feature importance", "Phát hiện mô hình xuống cấp theo thời gian"],
        ],
        "Bảng 6.8: Bộ chỉ số nên theo dõi khi triển khai",
    )
    add_paragraph(
        doc,
        "Một khuyến nghị quan trọng là luôn giữ seasonal naive làm benchmark vận hành. Seasonal naive đơn giản, dễ hiểu và rất cạnh tranh trong dữ liệu này. Nếu một phiên bản mô hình học máy mới không vượt được seasonal naive trên validation hoặc không cải thiện business metric, không nên triển khai chỉ vì mô hình phức tạp hơn.",
    )
    add_table(
        doc,
        ["Mức ưu tiên", "Điều kiện nhận diện", "Hành động đề xuất"],
        [
            ["Cao", "Stockout rate cao và recovered lift cao", "Tăng kiểm tra tồn kho, điều chỉnh safety stock, ưu tiên bổ sung"],
            ["Trung bình", "Seasonality tuần rõ nhưng stockout vừa phải", "Dùng seasonal forecast làm kế hoạch nền, cập nhật theo LightGBM"],
            ["Cảnh báo", "Velocity tăng nhanh trước stockout", "Kiểm tra tồn kho trong ngày, bổ sung nhanh nếu có thể"],
            ["Theo dõi", "Peer sales tăng khi sản phẩm chính stockout", "Đánh giá quan hệ thay thế và bố trí nhóm hàng"],
            ["Thận trọng", "Forecast interval rộng", "Không tự động nhập cao; cần xem chi phí tồn kho và hư hỏng"],
        ],
        "Bảng 6.9: Ma trận hành động vận hành dựa trên kết quả mô hình",
    )
    add_paragraph(
        doc,
        "Ngoài ra, hệ thống cần cơ chế governance. Các tham số như cap q90, calibration window, retraining frequency và ngưỡng cảnh báo không nên cố định vĩnh viễn. Chúng cần được theo dõi định kỳ vì hành vi mua hàng, chương trình khuyến mãi, danh mục sản phẩm và chính sách vận hành có thể thay đổi theo thời gian.",
    )
    add_paragraph(
        doc,
        "Tóm lại, giá trị thực tế của mô hình nằm ở việc giúp doanh nghiệp nhìn thấy phần nhu cầu bị che bởi stockout, giữ được nhịp mùa vụ tuần trong kế hoạch nhập hàng và lượng hóa rủi ro khi ra quyết định. Đây là bước chuyển từ dự báo như một bài toán chỉ số sang dự báo như một công cụ hỗ trợ vận hành.",
    )

    doc.add_heading("Kết luận Chương 6", level=2)
    add_paragraph(
        doc,
        "Chương 6 cho thấy các kết quả mô hình có thể được diễn giải thành insight kinh doanh cụ thể. Stockout làm observed sales đánh giá thấp demand; recovered demand giúp ước lượng lost sales; seasonality tuần cung cấp baseline vận hành mạnh; velocity và peer/substitution giúp nhận diện bối cảnh trước và trong stockout; còn forecast next-7-day hỗ trợ quy trình replenishment.",
    )
    add_paragraph(
        doc,
        "Điểm quan trọng nhất là mô hình không nên được triển khai như một hộp đen tự động nhập hàng. Nó nên được dùng như một hệ thống hỗ trợ quyết định có kiểm soát: luôn so với seasonal baseline, theo dõi forecast bias, kiểm tra stockout/lost sales sau triển khai và điều chỉnh theo chi phí thực tế của thiếu hàng và tồn kho dư.",
    )

    saved_path = save_doc(doc, OUT_PATH)
    print(saved_path)


if __name__ == "__main__":
    build_doc()
