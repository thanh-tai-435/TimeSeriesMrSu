from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TABLE_DIR = BASE_DIR / "outputs" / "tables"
FIGURE_DIR = BASE_DIR / "outputs" / "figures"
OUT_DIR = BASE_DIR / "deliverables" / "reports"
OUT_PATH = OUT_DIR / "Chuong_5_Chan_doan_mo_hinh_va_do_bat_dinh.docx"
IMAGE_OUT_DIR = OUT_DIR / "Chuong_5_Hinh_anh_goc"


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
        doc.save(fallback)
        return fallback


def set_default_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12 if style_name == "Normal" else 14)
        style.font.bold = style_name.startswith("Heading")
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)

    for style_name, italic, bold in [("Caption", True, False), ("Table Caption", False, True)]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.font.italic = italic
        style.font.bold = bold
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    doc.add_paragraph(caption, style="Table Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, export_filename: str, width: float = 6.2) -> None:
    path = FIGURE_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[Thiếu hình: {filename}]")
        return
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, IMAGE_OUT_DIR / export_filename)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")


def fmt_float(value: str | float, digits: int = 4) -> str:
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)


def fmt_pct(value: str | float, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def fmt_int(value: str | float) -> str:
    try:
        return f"{int(float(value)):,}"
    except Exception:
        return str(value)


def format_cell(col: str, value) -> str:
    if isinstance(value, float):
        lowered = col.lower()
        if "p-value" in lowered:
            return f"{value:.4f}"
        if any(key in lowered for key in ["wape", "smape", "wpe", "coverage", "weight"]):
            return fmt_pct(value)
        return fmt_float(value, 4)
    return str(value)


def ordered_models() -> list[str]:
    return [
        "Recovered naive x7",
        "Recovered seasonal naive 7-day",
        "Recovered rolling mean 14-day",
        "Observed-sales forecasting",
        "Recovered-demand forecasting",
        "Recovered seasonal-ML hybrid",
    ]


def table_rows(df: pd.DataFrame, cols: list[str]) -> list[list[str]]:
    return [[format_cell(col, row[col]) for col in cols] for _, row in df.iterrows()]


def diagnostics_df() -> pd.DataFrame:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_diagnostics.csv")
    df["rank"] = df["Model"].map({name: idx for idx, name in enumerate(ordered_models())})
    return df.sort_values("rank").reset_index(drop=True)


def nonoverlap_df() -> pd.DataFrame:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_nonoverlap_diagnostics.csv")
    df["rank"] = df["Model"].map({name: idx for idx, name in enumerate(ordered_models())})
    return df.sort_values("rank").reset_index(drop=True)


def interval_summary_df() -> pd.DataFrame:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_prediction_intervals.csv")
    rows = []
    for model_name in ordered_models():
        group = df[df["Model"] == model_name]
        if group.empty:
            continue
        coverage_80 = ((group["y_true_recovered"] >= group["lower_80"]) & (group["y_true_recovered"] <= group["upper_80"])).mean()
        coverage_95 = ((group["y_true_recovered"] >= group["lower_95"]) & (group["y_true_recovered"] <= group["upper_95"])).mean()
        rows.append(
            {
                "Model": model_name,
                "80% coverage": coverage_80,
                "95% coverage": coverage_95,
                "Mean 80% width": (group["upper_80"] - group["lower_80"]).mean(),
                "Mean 95% width": (group["upper_95"] - group["lower_95"]).mean(),
                "N": len(group),
            }
        )
    summary = pd.DataFrame(rows)
    summary.to_csv(TABLE_DIR / "owner_two_stage_interval_coverage_summary.csv", index=False)
    return summary


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if IMAGE_OUT_DIR.exists():
        shutil.rmtree(IMAGE_OUT_DIR)
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)

    diagnostics = diagnostics_df()
    nonoverlap = nonoverlap_df()
    intervals = interval_summary_df()
    blend = pd.read_csv(TABLE_DIR / "owner_hybrid_blend_selection.csv").head(5)

    hybrid = diagnostics[diagnostics["Model"] == "Recovered seasonal-ML hybrid"].iloc[0]
    seasonal = diagnostics[diagnostics["Model"] == "Recovered seasonal naive 7-day"].iloc[0]
    recovered_ml = diagnostics[diagnostics["Model"] == "Recovered-demand forecasting"].iloc[0]
    observed_ml = diagnostics[diagnostics["Model"] == "Observed-sales forecasting"].iloc[0]
    rel_gain_vs_observed = (observed_ml["WAPE"] - hybrid["WAPE"]) / observed_ml["WAPE"]
    rel_gain_vs_seasonal = (seasonal["WAPE"] - hybrid["WAPE"]) / seasonal["WAPE"]

    doc = Document()
    set_default_font(doc)

    title = doc.add_heading("CHƯƠNG 5: CHẨN ĐOÁN MÔ HÌNH VÀ ĐÁNH GIÁ ĐỘ BẤT ĐỊNH", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        "Các chương trước đã trình bày lý do cần khôi phục latent demand và kết quả so sánh giữa các mô hình dự báo. Tuy nhiên, một mô hình dự báo không nên được đánh giá chỉ bằng một con số WAPE. Với bài toán bán lẻ có stockout, chuỗi thưa và target next-7-day chồng lấn, cần kiểm tra thêm ba khía cạnh: phần dư còn chứa pattern hệ thống hay không, kết luận có ổn định khi loại bớt target window chồng lấn hay không, và forecast có đi kèm vùng bất định đủ hữu ích cho quyết định nhập hàng hay không.",
    )
    add_paragraph(
        doc,
        "Tinh thần của Chương 5 là kiểm tra tính đáng tin của kết quả, không phải cố chứng minh mô hình hoàn hảo. Nếu residual còn tự tương quan hoặc không chuẩn, điều đó cần được ghi nhận và giải thích đúng. Trong môi trường vận hành thật, một mô hình hữu ích là mô hình giúp giảm sai lệch ra quyết định, có benchmark rõ ràng, có kiểm tra độ ổn định và có cách lượng hóa rủi ro; không nhất thiết phải thỏa mãn tuyệt đối mọi giả định lý thuyết của chuỗi thời gian cổ điển.",
    )
    add_paragraph(
        doc,
        "Đối tượng chẩn đoán chính trong chương này là các dự báo trên recovered latent demand proxy. Đây là target đã được xử lý stockout ở cấp giờ và tổng hợp thành tổng nhu cầu 7 ngày tiếp theo. Vì vậy, các chỉ số trong chương này cần được đọc theo đúng framing của đồ án: mô hình đang dự báo nhu cầu đã giảm bias do stockout, chứ không chỉ dự báo doanh số quan sát bị giới hạn bởi tồn kho.",
    )

    doc.add_heading("5.1. Chẩn đoán Phần dư (Residual Diagnostics)", level=2)
    add_paragraph(
        doc,
        "Phần dư được định nghĩa là chênh lệch giữa recovered latent demand proxy và giá trị dự báo. Nếu phần dư có trung bình gần 0, mô hình ít bị bias tổng thể. Nếu WPE âm, dự báo có xu hướng cao hơn thực tế; nếu WPE dương, mô hình có xu hướng under-forecast. Bên cạnh đó, residual ACF và kiểm định Ljung-Box được dùng để kiểm tra phần dư còn tự tương quan hay không. Kiểm định Jarque-Bera được dùng để kiểm tra hình dạng phân phối phần dư có gần chuẩn hay không.",
    )
    add_paragraph(
        doc,
        "Trong bài toán này, không nên kỳ vọng phần dư hoàn toàn white noise. Lý do thứ nhất là target dự báo là tổng nhu cầu 7 ngày tiếp theo, nên forecast origin liên tiếp có phần target window chồng lấn. Lý do thứ hai là dữ liệu bán lẻ có nhiều chuỗi nhỏ, nhiều giá trị 0, spike theo khuyến mãi và biến động do stockout. Lý do thứ ba là recovered demand chỉ là proxy có kiểm soát, không phải ground truth tuyệt đối. Vì vậy, diagnostics cần được đọc theo hướng xác định rủi ro còn lại, không phải như tiêu chí loại bỏ máy móc.",
    )
    add_table(
        doc,
        ["Model", "RMSE", "MAE", "WAPE", "WPE", "Residual mean", "Residual std", "Ljung-Box p-value", "Jarque-Bera p-value", "N"],
        table_rows(
            diagnostics,
            ["Model", "RMSE", "MAE", "WAPE", "WPE", "Residual mean", "Residual std", "Ljung-Box p-value", "Jarque-Bera p-value", "N"],
        ),
        "Bảng 5.1: Chẩn đoán phần dư của các mô hình trên recovered latent demand proxy",
    )
    add_paragraph(
        doc,
        "Bảng 5.1 cho thấy mô hình hybrid seasonal-ML đạt WAPE khoảng "
        f"{fmt_pct(hybrid['WAPE'])}, thấp hơn observed-sales forecasting với WAPE khoảng {fmt_pct(observed_ml['WAPE'])}. "
        f"Mức cải thiện tương đối so với observed-sales forecasting là khoảng {fmt_pct(rel_gain_vs_observed)}. "
        "Điều này xác nhận kết luận chính của Chương 4: xử lý stockout trước khi forecast giúp giảm sai số khi mục tiêu là recovered demand. Đồng thời, hybrid chỉ cải thiện nhẹ so với seasonal naive 7-day, khoảng "
        f"{fmt_pct(rel_gain_vs_seasonal)}, cho thấy seasonal pattern là một baseline rất mạnh và không nên bị loại khỏi mô hình cuối.",
    )
    add_paragraph(
        doc,
        "WPE của hybrid vẫn âm, nghĩa là mô hình có xu hướng dự báo cao hơn recovered target ở mức tổng thể. Tuy nhiên, mức bias này nhỏ hơn nhiều so với observed-sales forecasting khi evaluated trên recovered demand. Đây là điểm quan trọng về mặt nghiệp vụ: mô hình forecast trực tiếp observed sales dễ under-estimate demand thật vì học từ dữ liệu bị stockout censor, trong khi mô hình two-stage giảm đáng kể sai lệch đó.",
    )
    add_figure(
        doc,
        "owner_two_stage_diagnostics_wape.png",
        "Hình 5.1: So sánh WAPE trong chẩn đoán two-stage",
        "Hinh_5.1_So_sanh_WAPE_diagnostics.png",
    )
    add_figure(
        doc,
        "owner_two_stage_bias_comparison.png",
        "Hình 5.2: So sánh xu hướng bias giữa các mô hình",
        "Hinh_5.2_So_sanh_bias_giua_cac_mo_hinh.png",
    )
    add_paragraph(
        doc,
        "Residual ACF trong Hình 5.3 cho thấy phần dư vẫn còn một số tương quan ở các lag ngắn. Kết quả này phù hợp với cấu trúc target chồng lấn và không làm vô hiệu mô hình. Thay vào đó, nó cho thấy còn tồn tại pattern ngắn hạn hoặc nhóm sản phẩm-cửa hàng mà mô hình chưa học hết. Đây là lý do đồ án không khẳng định mô hình đã tạo ra phần dư hoàn toàn độc lập, mà chỉ kết luận rằng mô hình đã cải thiện đáng kể so với hướng forecast trực tiếp observed sales.",
    )
    add_figure(
        doc,
        "owner_two_stage_residual_acf.png",
        "Hình 5.3: ACF phần dư của mô hình two-stage tốt nhất",
        "Hinh_5.3_ACF_phan_du_two_stage.png",
    )
    add_figure(
        doc,
        "owner_two_stage_residual_distribution.png",
        "Hình 5.4: Phân phối phần dư của mô hình two-stage tốt nhất",
        "Hinh_5.4_Phan_phoi_phan_du_two_stage.png",
    )
    add_paragraph(
        doc,
        "Phân phối phần dư trong Hình 5.4 không hoàn toàn chuẩn. Đây là đặc điểm thường gặp trong dữ liệu bán lẻ nhiều zero và nhiều spike. Với bài toán này, giả định chuẩn của phần dư không phải điều kiện trung tâm như trong một số mô hình thống kê cổ điển. Điều quan trọng hơn là phần dư có được theo dõi định kỳ hay không, có bias lớn theo nhóm hàng/cửa hàng hay không, và forecast có đủ ổn định để hỗ trợ replenishment hay không.",
    )
    add_paragraph(
        doc,
        "Kết quả residual diagnostics vì vậy được diễn giải theo hướng thực tế: mô hình hybrid seasonal-ML là lựa chọn tốt nhất trong các mô hình được thử, nhưng vẫn còn residual structure. Trong báo cáo và thuyết trình, điểm này cần được nói thẳng. Việc thừa nhận residual chưa white noise hoàn toàn làm cho kết luận đáng tin hơn, vì nó cho thấy mô hình được đánh giá bằng tiêu chuẩn kiểm định chứ không chỉ bằng chỉ số chính thuận lợi.",
    )

    doc.add_heading("5.2. Kiểm định Sự Bất tương trùng (Non-overlap Diagnostics)", level=2)
    add_paragraph(
        doc,
        "Một điểm kỹ thuật quan trọng của bài toán là target next-7-day được tạo theo forecast origin hằng ngày. Khi origin dịch một ngày, cửa sổ target 7 ngày mới vẫn chia sẻ 6 ngày với cửa sổ cũ. Điều này làm tăng số lượng mẫu đánh giá nhưng cũng làm các quan sát evaluation không hoàn toàn độc lập. Nếu chỉ nhìn metrics trên toàn bộ origin liên tiếp, kết quả có thể trông ổn định hơn thực tế.",
    )
    add_paragraph(
        doc,
        "Để kiểm tra vấn đề này, đồ án bổ sung non-overlap diagnostics. Thay vì dùng mọi forecast origin trong test set, ta lấy các origin cách nhau 7 ngày để các target window ít chồng lấn hơn. Đây không phải thay thế hoàn toàn evaluation chính, vì số mẫu sẽ giảm, nhưng là một kiểm tra độ bền cần thiết. Nếu thứ hạng mô hình thay đổi mạnh trong non-overlap setting, kết luận ở Chương 4 sẽ cần được xem xét lại.",
    )
    add_table(
        doc,
        ["Model", "WAPE", "WPE", "Residual mean", "Residual std", "Ljung-Box p-value", "N", "Aggregate residual points"],
        table_rows(
            nonoverlap,
            ["Model", "WAPE", "WPE", "Residual mean", "Residual std", "Ljung-Box p-value", "N", "Aggregate residual points"],
        ),
        "Bảng 5.2: Diagnostics trên các target window không chồng lấn",
    )
    hybrid_non = nonoverlap[nonoverlap["Model"] == "Recovered seasonal-ML hybrid"].iloc[0]
    seasonal_non = nonoverlap[nonoverlap["Model"] == "Recovered seasonal naive 7-day"].iloc[0]
    observed_non = nonoverlap[nonoverlap["Model"] == "Observed-sales forecasting"].iloc[0]
    add_paragraph(
        doc,
        "Bảng 5.2 cho thấy khi giảm target overlap, kết luận chính vẫn được giữ. Hybrid seasonal-ML đạt WAPE khoảng "
        f"{fmt_pct(hybrid_non['WAPE'])}, vẫn tốt hơn observed-sales forecasting với WAPE khoảng {fmt_pct(observed_non['WAPE'])}. "
        "Khoảng cách giữa hybrid và seasonal naive vẫn nhỏ, vì seasonal naive 7-day vốn là benchmark rất mạnh trong dữ liệu có weekly seasonality. Điều quan trọng là kết quả không bị đảo chiều: mô hình two-stage vẫn cho thấy lợi ích so với forecast trực tiếp observed sales.",
    )
    add_paragraph(
        doc,
        "Số lượng aggregate residual points trong non-overlap diagnostics thấp hơn nhiều so với đánh giá đầy đủ. Vì vậy, kiểm định thống kê trong bảng này không nên được đọc như bằng chứng tuyệt đối. Vai trò của nó là kiểm tra tính nhất quán của kết luận khi loại bớt sự phụ thuộc do target window chồng lấn. Việc hybrid vẫn giữ WAPE thấp nhất hoặc nằm sát seasonal naive cho thấy kết luận của đồ án không chỉ là sản phẩm của cách tạo nhiều origin chồng lấn.",
    )
    add_table(
        doc,
        ["LightGBM weight", "Seasonal naive weight", "Validation WAPE"],
        table_rows(blend, ["LightGBM weight", "Seasonal naive weight", "Validation WAPE"]),
        "Bảng 5.3: Độ nhạy của trọng số hybrid trên validation set",
    )
    add_paragraph(
        doc,
        "Bảng 5.3 bổ sung một kiểm tra độ ổn định khác: các trọng số hybrid tốt nhất trên validation nằm khá gần nhau. Cấu hình 0.60 LightGBM và 0.40 seasonal naive có WAPE validation thấp nhất, trong khi các cấu hình lân cận như 0.65/0.35 hoặc 0.55/0.45 cho kết quả gần tương đương. Điều này cho thấy kết luận hybrid không phụ thuộc quá nhạy vào một trọng số duy nhất. Về mặt nghiệp vụ, có thể hiểu rằng mô hình cần giữ một phần đáng kể seasonal baseline, đồng thời dùng LightGBM để hiệu chỉnh theo thông tin sản phẩm, cửa hàng, stockout và lịch sử demand.",
    )

    doc.add_heading("5.3. Phân tích Khoảng Dự báo (Prediction Intervals) và Định lượng Độ Cậy", level=2)
    add_paragraph(
        doc,
        "Trong vận hành nhập hàng, point forecast không đủ. Nếu chỉ có một giá trị dự báo trung bình, người ra quyết định không biết nên nhập sát forecast, nhập cao hơn để giảm rủi ro hết hàng, hay nhập thấp hơn để tránh hư hỏng hàng tươi sống. Vì vậy, khoảng dự báo được dùng để lượng hóa vùng bất định quanh point forecast. Khoảng này không nói nhu cầu chắc chắn nằm trong một biên, mà cung cấp một mức tham chiếu rủi ro dựa trên sai số validation.",
    )
    add_paragraph(
        doc,
        "Trong đồ án, prediction intervals được xây dựng theo cách thực dụng: lấy phân vị của absolute residual trên validation set, sau đó cộng/trừ quanh point forecast trên test set. Cách làm này không phức tạp như mô hình Bayesian hoặc quantile regression đầy đủ, nhưng có ba ưu điểm phù hợp với phạm vi đồ án: dễ tái lập, dễ giải thích và tách biệt giữa tập chọn khoảng với tập test cuối. Điều này cũng tránh việc điều chỉnh khoảng dự báo bằng thông tin từ test set.",
    )
    add_table(
        doc,
        ["Model", "80% coverage", "95% coverage", "Mean 80% width", "Mean 95% width", "N"],
        table_rows(intervals, ["Model", "80% coverage", "95% coverage", "Mean 80% width", "Mean 95% width", "N"]),
        "Bảng 5.4: Độ phủ và độ rộng trung bình của khoảng dự báo",
    )
    hybrid_interval = intervals[intervals["Model"] == "Recovered seasonal-ML hybrid"].iloc[0]
    add_paragraph(
        doc,
        "Với mô hình hybrid, khoảng 80% đạt coverage khoảng "
        f"{fmt_pct(hybrid_interval['80% coverage'])}, còn khoảng 95% đạt coverage khoảng {fmt_pct(hybrid_interval['95% coverage'])}. "
        "Coverage 80% thấp hơn mức danh nghĩa cho thấy khoảng 80% còn hơi hẹp so với biến động thật của dữ liệu. Ngược lại, khoảng 95% có coverage gần mức kỳ vọng hơn nhưng đổi lại độ rộng lớn hơn. Đây là trade-off tự nhiên giữa độ an toàn và chi phí tồn kho.",
    )
    add_figure(
        doc,
        "owner_two_stage_forecast_interval.png",
        "Hình 5.5: Khoảng dự báo của mô hình two-stage trên recovered latent demand",
        "Hinh_5.5_Khoang_du_bao_two_stage.png",
    )
    add_paragraph(
        doc,
        "Hình 5.5 minh họa cách khoảng dự báo bao quanh forecast aggregate. Khi nhu cầu thực tế biến động mạnh, khoảng dự báo rộng giúp người vận hành thấy rằng point forecast không nên được xem là con số chắc chắn. Với nhóm hàng có rủi ro stockout cao hoặc thời gian bổ sung hàng dài, doanh nghiệp có thể ra quyết định gần upper bound hơn. Với nhóm hàng dễ hư hỏng, quyết định có thể nghiêng về point forecast hoặc vùng thấp hơn, kèm tần suất cập nhật forecast cao hơn.",
    )
    add_table(
        doc,
        ["Mức sử dụng", "Ý nghĩa vận hành", "Khi nào nên dùng"],
        [
            ["Point forecast", "Kế hoạch nhập hàng trung bình", "Nhóm hàng ổn định, chi phí thiếu hàng và tồn kho cân bằng"],
            ["80% interval", "Vùng dao động thường gặp", "Theo dõi rủi ro hằng tuần, cảnh báo forecast bất thường"],
            ["95% interval", "Vùng an toàn cao hơn", "Nhóm dễ stockout, lead time dài, hoặc service level quan trọng"],
            ["Upper bound", "Kế hoạch thận trọng chống thiếu hàng", "Sự kiện, khuyến mãi, nhóm hàng chiến lược"],
        ],
        "Bảng 5.5: Diễn giải khoảng dự báo trong quyết định replenishment",
    )
    add_paragraph(
        doc,
        "Điểm quan trọng là prediction interval không thay thế quyết định nghiệp vụ. Nó giúp lượng hóa bất định để người vận hành chọn mức nhập phù hợp với mục tiêu cụ thể. Nếu mục tiêu là giảm stockout, có thể dùng forecast cao hơn point estimate. Nếu mục tiêu là giảm hủy hàng tươi sống, có thể dùng mức thận trọng hơn và tăng tần suất cập nhật. Vì vậy, phần uncertainty trong đồ án gắn trực tiếp với bài toán kinh doanh, không chỉ là một biểu đồ bổ sung.",
    )

    doc.add_heading("Kết luận Chương 5", level=2)
    add_paragraph(
        doc,
        "Chương 5 cho thấy mô hình hybrid seasonal-ML không chỉ có WAPE tốt nhất trong nhóm mô hình chính mà còn giữ được kết luận khi kiểm tra non-overlap diagnostics. Tuy nhiên, residual chưa hoàn toàn white noise và phân phối phần dư chưa chuẩn. Đây là hạn chế cần được trình bày trung thực, đồng thời là cơ sở cho hướng phát triển tiếp theo như thêm feature theo nhóm hàng, mô hình hóa substitution sâu hơn hoặc dùng probabilistic forecasting.",
    )
    add_paragraph(
        doc,
        "Về mặt vận hành, đóng góp chính của chương này là chuyển forecast từ một con số đơn lẻ sang một quyết định có nhận thức rủi ro. Point forecast cho biết kỳ vọng nhu cầu, residual diagnostics cho biết mô hình còn sai ở đâu, non-overlap diagnostics kiểm tra độ bền của kết luận, còn prediction intervals giúp định lượng vùng bất định để ra quyết định nhập hàng. Nhờ vậy, mô hình không chỉ được đánh giá bằng độ chính xác, mà còn bằng khả năng hỗ trợ quyết định trong điều kiện dữ liệu bán lẻ nhiều nhiễu và stockout.",
    )

    saved_path = save_doc(doc, OUT_PATH)
    print(saved_path)


if __name__ == "__main__":
    build_doc()
