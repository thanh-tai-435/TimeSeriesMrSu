from __future__ import annotations

from pathlib import Path
import csv

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TABLE_DIR = BASE_DIR / "outputs" / "tables"
OUT_DIR = BASE_DIR / "deliverables" / "reports"
OUT_PATH = OUT_DIR / "Chuong_7_Tong_ket_va_huong_phat_trien.docx"


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
        doc.save(fallback)
        return fallback


def set_default_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12 if style_name == "Normal" else 14)
        style.font.bold = style_name.startswith("Heading")
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)

    if "Table Caption" not in doc.styles:
        doc.styles.add_style("Table Caption", WD_STYLE_TYPE.PARAGRAPH)
    table_caption = doc.styles["Table Caption"]
    table_caption.font.name = "Times New Roman"
    table_caption.font.size = Pt(11)
    table_caption.font.bold = True
    table_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table_caption.paragraph_format.line_spacing = 1.5
    table_caption.paragraph_format.space_before = Pt(0)
    table_caption.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    doc.add_paragraph(caption, style="Table Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for paragraph in cells[idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    doc.add_paragraph()


def read_metric_table(filename: str) -> dict[str, str]:
    result: dict[str, str] = {}
    with (TABLE_DIR / filename).open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            metric_key = "Metric" if "Metric" in row else "metric"
            value_key = "Value" if "Value" in row else "value"
            result[row[metric_key]] = row[value_key]
    return result


def fmt_float(value: str | float, digits: int = 2) -> str:
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)


def fmt_pct(value: str | float, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def forecast_summary_rows() -> tuple[list[list[str]], dict[str, float]]:
    df = pd.read_csv(TABLE_DIR / "owner_two_stage_forecasting_comparison.csv")
    df = df[df["Evaluation target"] == "Recovered latent demand proxy"].copy()
    order = [
        "Recovered naive x7",
        "Recovered seasonal naive 7-day",
        "Recovered rolling mean 14-day",
        "Observed-sales forecasting",
        "Recovered-demand forecasting",
        "Recovered seasonal-ML hybrid",
    ]
    df["rank"] = df["Model"].map({name: idx for idx, name in enumerate(order)})
    df = df.sort_values("rank")

    rows: list[list[str]] = []
    metrics: dict[str, float] = {}
    for _, row in df.iterrows():
        rows.append(
            [
                row["Model"],
                fmt_pct(row["WAPE"]),
                fmt_pct(row["WPE"]),
                row["Training target"],
            ]
        )
        metrics[f"{row['Model']} WAPE"] = float(row["WAPE"])
        metrics[f"{row['Model']} WPE"] = float(row["WPE"])
    return rows, metrics


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    recovery = read_metric_table("owner_latent_recovery_summary.csv")
    data_quality = read_metric_table("data_quality_summary.csv")
    pseudo = pd.read_csv(TABLE_DIR / "imputation_pseudo_stockout_aggregate_validation.csv")
    interval = pd.read_csv(TABLE_DIR / "owner_two_stage_interval_coverage_summary.csv")
    diag = pd.read_csv(TABLE_DIR / "owner_two_stage_diagnostics.csv")
    forecast_rows, metrics = forecast_summary_rows()

    observed_wape = metrics["Observed-sales forecasting WAPE"]
    recovered_wape = metrics["Recovered-demand forecasting WAPE"]
    seasonal_wape = metrics["Recovered seasonal naive 7-day WAPE"]
    hybrid_wape = metrics["Recovered seasonal-ML hybrid WAPE"]
    hybrid_wpe = metrics["Recovered seasonal-ML hybrid WPE"]
    rel_gain = (observed_wape - hybrid_wape) / observed_wape

    pseudo_daily = pseudo[pseudo["Validation level"] == "Daily aggregate"].iloc[0]
    hybrid_interval = interval[interval["Model"] == "Recovered seasonal-ML hybrid"].iloc[0]
    hybrid_diag = diag[diag["Model"] == "Recovered seasonal-ML hybrid"].iloc[0]

    doc = Document()
    set_default_font(doc)

    title = doc.add_heading("CHƯƠNG 7: TỔNG KẾT VÀ HƯỚNG PHÁT TRIỂN", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        "Chương cuối không lặp lại toàn bộ kết quả đã trình bày ở các chương trước, mà đóng vai trò tổng hợp luận điểm nghiên cứu của đề tài. Trọng tâm của đồ án là chứng minh rằng bài toán dự báo bán lẻ hàng tươi sống không nên được nhìn như một bài toán forecasting thuần túy trên doanh số quan sát. Khi xảy ra hết hàng, doanh số quan sát chỉ phản ánh lượng hàng có thể bán được, không phản ánh đầy đủ nhu cầu thị trường. Vì vậy, đóng góp chính của đề tài nằm ở cách tái định nghĩa mục tiêu dự báo: từ dự báo observed sales sang dự báo latent demand đã được khôi phục có kiểm soát.",
    )
    add_paragraph(
        doc,
        "Toàn bộ pipeline được xây dựng theo logic hai giai đoạn. Giai đoạn thứ nhất xử lý stockout ở cấp độ giờ bằng expanding-window recovery để hạn chế rò rỉ dữ liệu tương lai. Giai đoạn thứ hai tổng hợp nhu cầu đã khôi phục lên cấp độ ngày và xây dựng mô hình dự báo phục vụ bài toán nhập hàng. Cách tiếp cận này giúp liên kết phần phân tích chuỗi thời gian với nhu cầu vận hành thực tế: phát hiện mất doanh số do hết hàng, ước lượng nhu cầu tiềm ẩn, sau đó dự báo nhu cầu tương lai để hỗ trợ replenishment.",
    )

    doc.add_heading("7.1. Kết luận Chung", level=2)
    add_paragraph(
        doc,
        "Về mặt dữ liệu, đề tài sử dụng tập FreshRetailNet-50K được lấy mẫu theo chuỗi với tỷ lệ "
        f"{fmt_pct(data_quality.get('Sample fraction', '0.1'))}. Cách lấy mẫu này giữ nguyên trục thời gian của từng store-product series, giúp các phân tích mùa vụ, stockout và split train-validation-test không bị phá vỡ. Sau khi tiền xử lý, dữ liệu hourly sample có "
        f"{fmt_float(data_quality.get('Hourly rows', ''), 0)} dòng, bao phủ {fmt_float(data_quality.get('Series', ''), 0)} chuỗi, "
        f"{fmt_float(data_quality.get('Stores', ''), 0)} cửa hàng và {fmt_float(data_quality.get('Products', ''), 0)} sản phẩm. Tỷ lệ stockout ở cấp độ dòng đạt khoảng "
        f"{fmt_pct(recovery.get('Stockout row rate', ''))}, đủ lớn để ảnh hưởng đáng kể đến target dự báo nếu chỉ dùng observed sales.",
    )
    add_paragraph(
        doc,
        "Về mặt khôi phục nhu cầu, recovered latent demand đạt "
        f"{fmt_float(recovery.get('Recovered latent demand', ''), 2)}, cao hơn observed sales "
        f"{fmt_float(recovery.get('Observed sales', ''), 2)}. Phần recovered lost demand đạt "
        f"{fmt_float(recovery.get('Recovered lost demand', ''), 2)}, tương ứng mức lift khoảng "
        f"{fmt_pct(recovery.get('Recovered lift over observed', ''))}. Kết quả này cho thấy doanh số quan sát đang đánh giá thấp nhu cầu thực tế trong các giai đoạn hết hàng. Tuy nhiên, đề tài không xem recovered demand là ground truth tuyệt đối, mà xem đây là proxy có kiểm soát thông qua calibration, capping và pseudo-stockout validation.",
    )
    add_paragraph(
        doc,
        "Về mặt dự báo, mô hình dự báo trực tiếp trên observed sales đạt WAPE khoảng "
        f"{fmt_pct(observed_wape)} khi đánh giá trên recovered latent demand proxy. Khi chuyển sang target recovered demand, WAPE giảm còn "
        f"{fmt_pct(recovered_wape)}. Mô hình hybrid seasonal-ML đạt WAPE tốt nhất, khoảng {fmt_pct(hybrid_wape)}, tương ứng cải thiện tương đối khoảng {fmt_pct(rel_gain)} so với mô hình học trên observed sales. Điều này củng cố kết luận rằng xử lý stockout trước khi forecast không chỉ hợp lý về mặt lý thuyết mà còn tạo cải thiện thực nghiệm.",
    )
    add_paragraph(
        doc,
        "Một điểm đáng chú ý là Seasonal Naive 7-day đạt WAPE khoảng "
        f"{fmt_pct(seasonal_wape)}, rất gần với hybrid. Điều này không phải là điểm yếu của đồ án, mà là một phát hiện quan trọng: dữ liệu bán lẻ hàng tươi sống có mùa vụ tuần rất mạnh. Vì vậy, một mô hình học máy chỉ có ý nghĩa khi vượt qua hoặc bổ sung được baseline mùa vụ. Đề tài giữ Seasonal Naive như baseline trọng yếu và dùng LightGBM theo hướng hiệu chỉnh phi tuyến, không trình bày học máy như lựa chọn mặc định thay thế mọi baseline truyền thống.",
    )
    add_table(
        doc,
        ["Model", "WAPE", "WPE", "Training target"],
        forecast_rows,
        "Bảng 7.1: Tổng hợp hiệu suất các mô hình trên recovered latent demand proxy",
    )

    add_paragraph(
        doc,
        "Xét về tổng thể, đóng góp của đề tài có thể tóm gọn ở ba tầng. Tầng thứ nhất là framing đúng bài toán: observed sales bị censoring bởi stockout nên không đồng nhất với true demand. Tầng thứ hai là xây dựng cơ chế recovery có kiểm soát theo thời gian: expanding-window, warm-up, calibration và capping. Tầng thứ ba là nối recovery với forecasting và business insight: mô hình không chỉ xuất ra chỉ số WAPE, mà còn giúp ước lượng lost sales, đánh giá rủi ro hết hàng và hỗ trợ quy trình nhập hàng.",
    )
    add_table(
        doc,
        ["Tầng đóng góp", "Nội dung chính", "Ý nghĩa đối với đề tài"],
        [
            [
                "Framing bài toán",
                "Phân biệt observed sales và latent demand trong bối cảnh stockout",
                "Giúp đề tài tránh cách tiếp cận forecast trực tiếp trên dữ liệu bị kiểm duyệt",
            ],
            [
                "Latent demand recovery",
                "Khôi phục nhu cầu ở cấp độ giờ bằng expanding-window recovery, calibration và capping",
                "Giảm rủi ro leakage và hạn chế over-imputation",
            ],
            [
                "Daily forecasting",
                "Dự báo nhu cầu ngày bằng Seasonal Naive, LightGBM và hybrid seasonal-ML",
                "Bám sát bài toán nhập hàng theo ngày và giữ baseline mùa vụ mạnh",
            ],
            [
                "Diagnostics",
                "Kiểm tra residual, non-overlap và prediction intervals",
                "Không chỉ báo cáo điểm số mà còn đánh giá độ tin cậy của mô hình",
            ],
            [
                "Business insight",
                "Diễn giải lost sales, stockout risk và replenishment",
                "Chuyển kết quả kỹ thuật thành khuyến nghị vận hành",
            ],
        ],
        "Bảng 7.2: Các tầng đóng góp chính của đề tài",
    )

    add_paragraph(
        doc,
        "Do đó, kết luận chung của đồ án là: trong dữ liệu bán lẻ có stockout, việc dự báo trực tiếp doanh số quan sát có thể dẫn đến đánh giá thấp nhu cầu và tạo quyết định nhập hàng thiếu. Một pipeline hai giai đoạn, trong đó latent demand được khôi phục trước rồi mới dùng làm target dự báo, là hướng tiếp cận phù hợp hơn. Mô hình cuối cùng chưa phải là hệ thống tối ưu tuyệt đối, nhưng đã đủ chặt chẽ để thể hiện đầy đủ quy trình nghiên cứu: phát hiện vấn đề dữ liệu, xây dựng phương pháp, kiểm định chất lượng recovery, so sánh baseline, chẩn đoán mô hình và rút ra hàm ý kinh doanh.",
    )

    doc.add_heading("7.2. Các Hạn chế Tồn đọng của Đề tài", level=2)
    add_paragraph(
        doc,
        "Hạn chế đầu tiên là recovered demand vẫn chỉ là proxy. Trong các giờ stockout, nhu cầu thật không được quan sát trực tiếp. Mọi phương pháp khôi phục đều phải dựa trên giả định về hành vi bán trước đó, mùa vụ, peer demand và các tín hiệu lịch sử. Đề tài đã giảm rủi ro bằng expanding-window recovery, chỉ dùng dữ liệu quá khứ trong train-period, giữ 14 ngày warm-up và hiệu chỉnh bằng calibration factor. Tuy vậy, không thể khẳng định recovered demand là true demand tuyệt đối.",
    )
    add_paragraph(
        doc,
        "Hạn chế thứ hai là pseudo-stockout validation vẫn cho thấy xu hướng dự báo cao hơn thực tế. Ở cấp daily aggregate, WAPE của pseudo-stockout validation khoảng "
        f"{fmt_pct(pseudo_daily['WAPE'])}, trong khi prediction/actual ratio khoảng {fmt_float(pseudo_daily['Prediction / actual ratio'], 4)}. Điều này cho thấy imputation có thể làm tăng tổng sales nếu không được kiểm soát. Vì vậy, phần recovered lost demand cần được trình bày như một ước lượng có uncertainty, không phải con số doanh thu mất đi chắc chắn.",
    )
    add_paragraph(
        doc,
        "Hạn chế thứ ba nằm ở phần residual diagnostics. Với hybrid seasonal-ML, WPE khoảng "
        f"{fmt_pct(hybrid_wpe)} và Ljung-Box p-value ở lag 7 khoảng {fmt_float(hybrid_diag['Ljung-Box p-value'], 4)}. Kết quả này cho thấy mô hình đã giảm sai lệch đáng kể so với forecast trên observed sales, nhưng phần dư chưa hoàn toàn độc lập như white noise. Đây là đặc điểm thường gặp trong dữ liệu retail có nhiều zero, spike, sự kiện cục bộ và target next-7-day có tính chồng lấn giữa các forecast origin.",
    )
    add_paragraph(
        doc,
        "Hạn chế thứ tư là mô hình substitution mới chỉ được đưa vào ở mức feature engineering ban đầu. Đề tài đã bổ sung tín hiệu peer sales, nhóm store-category và velocity/momentum trước stockout để xử lý hai điểm mù quan trọng: sản phẩm A hết hàng làm sản phẩm B tăng bất thường, và tốc độ bán 16-18h phản ánh demand tăng mạnh trước khi missing từ 18h trở đi. Tuy nhiên, quan hệ thay thế sản phẩm trong thực tế còn phụ thuộc giá, khuyến mãi, vị trí trưng bày, giỏ hàng và hành vi khách hàng. Các yếu tố này chưa được mô hình hóa đầy đủ.",
    )
    add_paragraph(
        doc,
        "Hạn chế thứ năm là đề tài sử dụng sample 10% thay vì toàn bộ dữ liệu. Lựa chọn này hợp lý trong bối cảnh tài nguyên tính toán giới hạn và mục tiêu đồ án không yêu cầu full data, nhưng nó có thể làm giảm khả năng quan sát đầy đủ các quan hệ peer/substitution trong mạng sản phẩm-cửa hàng. Nếu triển khai thực tế, cần đánh giá lại trên toàn bộ dữ liệu hoặc ít nhất trên sample được thiết kế theo nhóm hàng, thành phố và mức độ stockout.",
    )
    add_table(
        doc,
        ["Hạn chế", "Ảnh hưởng", "Cách đã kiểm soát trong đề tài"],
        [
            [
                "Recovered demand chỉ là proxy",
                "Không quan sát trực tiếp nhu cầu thật trong giờ hết hàng",
                "Expanding-window, warm-up, calibration, capping",
            ],
            [
                "Imputation có thể over-predict",
                "Lost sales proxy có thể bị đánh giá cao",
                "Pseudo-stockout validation và sensitivity analysis theo cap",
            ],
            [
                "Residual chưa hoàn toàn white noise",
                "Còn pattern chưa được mô hình học hết",
                "Residual diagnostics, non-overlap diagnostics và prediction intervals",
            ],
            [
                "Substitution modeling còn đơn giản",
                "Chưa mô hình hóa đầy đủ quan hệ thay thế giữa sản phẩm",
                "Peer/category features và velocity trước stockout",
            ],
            [
                "Sample 10%",
                "Có thể thiếu một số quan hệ mạng sản phẩm-cửa hàng",
                "Lấy mẫu theo chuỗi để giữ nguyên cấu trúc thời gian",
            ],
        ],
        "Bảng 7.3: Hạn chế chính và cách kiểm soát trong phạm vi đồ án",
    )

    add_paragraph(
        doc,
        "Những hạn chế này không làm mất giá trị của đề tài, nhưng đặt ra ranh giới diễn giải rõ ràng. Kết quả tốt nhất không nên được hiểu là mô hình đã giải quyết hoàn toàn bài toán demand forecasting cho hàng tươi sống, mà nên được hiểu là một framework nghiên cứu có kiểm soát, đủ để chứng minh stockout-aware forecasting có giá trị hơn forecasting trực tiếp trên observed sales.",
    )

    doc.add_heading("7.3. Đề xuất Hướng Nghiên cứu và Phát triển Tương lai", level=2)
    add_paragraph(
        doc,
        "Hướng phát triển đầu tiên là mở rộng latent demand recovery bằng các mô hình chuyên cho missing hoặc censored time series. Các mô hình như SAITS, CSDI, TimesNet, ImputeFormer, GPVAE, iTransformer hoặc DLinear có thể được thử nghiệm trong tương lai. Tuy nhiên, bài học quan trọng từ đề tài là mô hình phức tạp chỉ nên được dùng khi đi kèm validation chặt chẽ. Nếu không kiểm soát leakage và over-imputation, deep learning có thể tạo ra kết quả có vẻ tốt nhưng khó tin cậy trong vận hành.",
    )
    add_paragraph(
        doc,
        "Hướng thứ hai là xây dựng substitution graph giữa các sản phẩm. Thay vì chỉ dùng peer sales theo store-category, có thể học quan hệ thay thế dựa trên tương quan nhu cầu, similarity danh mục, giá bán, khuyến mãi hoặc co-movement khi một sản phẩm stockout. Khi đó, nếu rau cải A hết hàng và rau cải B tăng bất thường, mô hình có thể phân biệt tốt hơn phần tăng do substitution với phần tăng do demand chung của cửa hàng.",
    )
    add_paragraph(
        doc,
        "Hướng thứ ba là cải thiện mô hình uncertainty. Hiện tại, khoảng dự báo của hybrid có 95% coverage khoảng "
        f"{fmt_pct(hybrid_interval['95% coverage'])} và 80% coverage khoảng {fmt_pct(hybrid_interval['80% coverage'])}. Đây là điểm khởi đầu tốt, nhưng để phục vụ nhập hàng, mô hình cần dự báo phân vị hoặc conformal prediction ổn định hơn. Điều này cho phép chuyển forecast thành quyết định safety stock, thay vì chỉ đưa ra một giá trị trung bình.",
    )
    add_paragraph(
        doc,
        "Hướng thứ tư là đánh giá mô hình bằng business metrics. Trong môi trường bán lẻ, một mô hình WAPE thấp chưa chắc là mô hình tốt nhất nếu nó làm tăng hàng tồn, hủy hàng hoặc chi phí vận hành. Do đó, các thí nghiệm tiếp theo nên đo thêm stockout rate sau khuyến nghị nhập hàng, service level, waste rate, inventory turnover và lợi nhuận ròng. Đây là bước cần thiết để chuyển từ mô hình học thuật sang công cụ hỗ trợ quyết định.",
    )
    add_paragraph(
        doc,
        "Hướng thứ năm là triển khai rolling retraining và drift monitoring. Nhu cầu hàng tươi sống có thể thay đổi theo mùa, thời tiết, ngày lễ, khuyến mãi, thay đổi giá và hành vi khách hàng. Vì vậy, mô hình cần dashboard theo dõi sai số theo tuần, store, category và nhóm stockout. Seasonal Naive nên tiếp tục được giữ làm benchmark thường trực: nếu mô hình học máy không còn vượt được baseline mùa vụ, hệ thống cần cảnh báo hoặc tái huấn luyện.",
    )
    add_table(
        doc,
        ["Hướng phát triển", "Mục tiêu", "Ưu tiên"],
        [
            [
                "Substitution graph",
                "Mô hình hóa quan hệ thay thế giữa sản phẩm khi stockout",
                "Cao, vì trực tiếp xử lý điểm mù business",
            ],
            [
                "Probabilistic forecasting",
                "Tạo quantile và prediction interval phục vụ safety stock",
                "Cao, vì cần cho replenishment thực tế",
            ],
            [
                "Deep imputation",
                "Cải thiện khôi phục demand trong các block stockout dài",
                "Trung bình, chỉ triển khai khi có validation chống leakage",
            ],
            [
                "Business metric evaluation",
                "Đo tác động đến stockout, waste và lợi nhuận",
                "Cao, nếu có dữ liệu vận hành bổ sung",
            ],
            [
                "Full-data retraining",
                "Kiểm tra độ ổn định khi mở rộng từ sample 10% lên toàn bộ dữ liệu",
                "Trung bình-cao, phụ thuộc tài nguyên tính toán",
            ],
            [
                "Drift monitoring",
                "Theo dõi mô hình sau triển khai và phát hiện thay đổi hành vi nhu cầu",
                "Cao trong giai đoạn production",
            ],
        ],
        "Bảng 7.4: Đề xuất hướng phát triển và mức độ ưu tiên",
    )

    add_paragraph(
        doc,
        "Nếu tiếp tục phát triển đề tài theo hướng ứng dụng, thứ tự ưu tiên hợp lý không phải là đưa ngay mô hình deep learning phức tạp vào pipeline. Ưu tiên nên là: kiểm định recovery tốt hơn, mô hình hóa substitution rõ hơn, lượng hóa uncertainty tốt hơn và đánh giá bằng business metrics. Sau khi các lớp kiểm định này ổn định, các mô hình sâu mới có ý nghĩa như một bước nâng cấp kỹ thuật.",
    )
    add_paragraph(
        doc,
        "Tóm lại, đề tài đã xây dựng được một hướng tiếp cận nhất quán cho bài toán dự báo chuỗi thời gian bán lẻ trong điều kiện stockout. Thay vì chỉ tối ưu điểm số dự báo, đề tài đặt trọng tâm vào việc hiểu đúng dữ liệu, khôi phục nhu cầu bị che khuất, kiểm soát rủi ro imputation và diễn giải kết quả thành quyết định vận hành. Đây là nền tảng phù hợp để mở rộng sang các mô hình recovery và forecasting nâng cao hơn trong tương lai.",
    )

    saved_path = save_doc(doc, OUT_PATH)
    print(saved_path)


if __name__ == "__main__":
    build_doc()
