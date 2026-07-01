from __future__ import annotations

import csv
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE_DIR = Path(__file__).resolve().parent
TABLE_DIR = BASE_DIR / "outputs" / "tables"
FIGURE_DIR = BASE_DIR / "outputs" / "figures"
OUT_DIR = BASE_DIR / "deliverables" / "reports"
OUT_PATH = OUT_DIR / "Chuong_2_EDA_Dac_tinh_chuoi_thoi_gian.docx"
IMAGE_OUT_DIR = OUT_DIR / "Chuong_2_Hinh_anh_goc"


def save_doc(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
        doc.save(fallback)
        return fallback


def set_default_font(doc: Document) -> None:
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12 if style_name == "Normal" else 14)
        style.font.bold = style_name.startswith("Heading")
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)

    for style_name, italic, bold in [("Caption", True, False), ("Table Caption", False, True)]:
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.font.italic = italic
        style.font.bold = bold
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Inches(0.25)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    doc.add_paragraph(caption, style="Table Caption")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    doc.add_paragraph()


def add_figure(doc: Document, filename: str, caption: str, export_filename: str, width: float = 6.2) -> None:
    path = FIGURE_DIR / filename
    if not path.exists():
        add_paragraph(doc, f"[Thiếu hình: {filename}]")
        return
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(path, IMAGE_OUT_DIR / export_filename)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")


def read_metric_table(filename: str) -> dict[str, str]:
    path = TABLE_DIR / filename
    result: dict[str, str] = {}
    with path.open("r", encoding="utf-8-sig", newline="") as file:
        reader = csv.DictReader(file)
        for row in reader:
            metric_key = "Metric" if "Metric" in row else "metric"
            value_key = "Value" if "Value" in row else "value"
            result[row[metric_key]] = row[value_key]
    return result


def fmt_float(value: str | float, digits: int = 4) -> str:
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return str(value)


def fmt_pct(value: str | float, digits: int = 2) -> str:
    try:
        return f"{float(value) * 100:.{digits}f}%"
    except Exception:
        return str(value)


def fmt_int(value: str | float) -> str:
    try:
        return f"{int(float(value)):,}"
    except Exception:
        return str(value)


def format_cell(col: str, value) -> str:
    if isinstance(value, float):
        if "p-value" in col.lower():
            return f"{value:.2e}"
        if any(key in col.lower() for key in ["rate", "wape", "smape", "lift", "share"]):
            return fmt_pct(value)
        return fmt_float(value, 4)
    return str(value)


def df_rows(filename: str, columns: list[str], max_rows: int | None = None) -> list[list[str]]:
    path = TABLE_DIR / filename
    if not path.exists():
        return []
    df = pd.read_csv(path)
    if max_rows is not None:
        df = df.head(max_rows)
    return [[format_cell(col, row[col]) for col in columns] for _, row in df.iterrows()]


def add_stationarity_table(doc: Document) -> None:
    rows = df_rows(
        "stationarity_tests.csv",
        ["Series", "Transformation", "ADF p-value", "KPSS p-value", "Conclusion"],
    )
    if rows:
        add_table(
            doc,
            ["Chuỗi", "Biến đổi", "ADF p-value", "KPSS p-value", "Kết luận"],
            rows,
            "Bảng 2.3: Tóm tắt kiểm định tính dừng bằng ADF và KPSS",
        )


def add_spectrum_table(doc: Document, filename: str, caption: str, unit_label: str) -> None:
    path = TABLE_DIR / filename
    if not path.exists():
        return
    df = pd.read_csv(path).head(5)
    period_col = "Period" if "Period" in df.columns else [col for col in df.columns if col.lower().startswith("period")][0]
    power_col = "Power" if "Power" in df.columns else "power"
    rows = [[fmt_float(row[period_col], 2), fmt_float(row[power_col], 4)] for _, row in df.iterrows()]
    add_table(doc, [f"Chu kỳ ({unit_label})", "Công suất phổ"], rows, caption)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if IMAGE_OUT_DIR.exists():
        shutil.rmtree(IMAGE_OUT_DIR)
    IMAGE_OUT_DIR.mkdir(parents=True, exist_ok=True)

    quality = read_metric_table("data_quality_summary.csv")
    sub_velocity_path = TABLE_DIR / "stockout_substitution_velocity_diagnostics.csv"
    sub_velocity_rows = (
        df_rows(
            "stockout_substitution_velocity_diagnostics.csv",
            [
                "Segment",
                "Rows",
                "Mean sale velocity ratio 3h/24h",
                "Mean sale momentum 3h-6h",
                "Mean peer sales same group",
                "Mean peer velocity ratio 3h/24h",
                "Mean peer stockout rate",
            ],
        )
        if sub_velocity_path.exists()
        else []
    )

    doc = Document()
    set_default_font(doc)

    title = doc.add_heading(
        "CHƯƠNG 2: PHÂN TÍCH KHÁM PHÁ DỮ LIỆU (EDA) VÀ ĐẶC TÍNH CHUỖI THỜI GIAN",
        level=1,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(
        doc,
        "Chương này phân tích dữ liệu FreshRetailNet-50K trước khi xây dựng mô hình dự báo. Trọng tâm không chỉ là mô tả dữ liệu, mà là xác định những đặc tính ảnh hưởng trực tiếp đến thiết kế bài toán: observed sales có thể bị kiểm duyệt bởi stockout, dữ liệu bán lẻ có nhiều giá trị 0, seasonality theo giờ và theo tuần khá rõ, đồng thời các chuỗi cửa hàng-sản phẩm có mức độ dị biệt lớn.",
    )
    add_paragraph(
        doc,
        "Trong phạm vi đồ án, dữ liệu được lấy mẫu 10% theo chuỗi. Cách lấy mẫu này giữ nguyên toàn bộ trục thời gian của các chuỗi được chọn, nhờ vậy vẫn bảo toàn seasonality, stockout và quan hệ trước-sau theo thời gian. Tập dữ liệu sau lấy mẫu gồm "
        f"{fmt_int(quality.get('Number of rows', ''))} quan sát theo giờ, "
        f"{fmt_int(quality.get('Number of series', ''))} chuỗi cửa hàng-sản phẩm, "
        f"{fmt_int(quality.get('Number of stores', ''))} cửa hàng, "
        f"{fmt_int(quality.get('Number of products', ''))} sản phẩm và "
        f"{fmt_int(quality.get('Number of cities', ''))} thành phố, trong giai đoạn từ "
        f"{quality.get('Start date', '')} đến {quality.get('End date', '')}.",
    )
    add_paragraph(
        doc,
        "Mạch phân tích của chương được triển khai theo logic phục vụ mô hình hóa. Mục 2.1 làm rõ vì sao observed sales không nên được xem ngay là nhu cầu thực. Mục 2.2 phân tích phân phối và mức độ thưa của doanh số. Mục 2.3, 2.5 và 2.6 kiểm tra seasonality, tính dừng, tự tương quan và phổ tần số để làm cơ sở cho baseline mùa vụ và lag features. Mục 2.4 xem xét các chuỗi đại diện để cho thấy bài toán không đồng nhất giữa các nhóm sản phẩm-cửa hàng. Các kết quả này dẫn trực tiếp sang phương pháp two-stage ở Chương 3.",
    )

    doc.add_heading("2.1. Phân tích Tổng Doanh số và Mối tương quan với Hiện tượng Hết hàng (Stockout)", level=2)
    add_paragraph(
        doc,
        "Bước đầu tiên của EDA là kiểm tra mức độ tin cậy của doanh số quan sát trước khi dùng nó làm target. Với dữ liệu bán lẻ tươi sống, doanh số ghi nhận không luôn phản ánh đầy đủ nhu cầu thực tế. Khi một sản phẩm hết hàng, lượng bán có thể bằng 0 hoặc thấp bất thường, nhưng điều này không đồng nghĩa với việc khách hàng không có nhu cầu. Vì vậy, tổng doanh số cần được đọc cùng với stockout rate.",
    )
    add_table(
        doc,
        ["Chỉ tiêu", "Giá trị"],
        [
            ["Số quan sát theo giờ", fmt_int(quality.get("Number of rows", ""))],
            ["Số chuỗi cửa hàng-sản phẩm", fmt_int(quality.get("Number of series", ""))],
            ["Tỷ lệ stockout", fmt_pct(quality.get("Stockout rate", ""))],
            ["Doanh số trung bình mỗi quan sát", fmt_float(quality.get("Average sale_amount", ""), 4)],
            ["Trung vị doanh số mỗi quan sát", fmt_float(quality.get("Median sale_amount", ""), 4)],
        ],
        "Bảng 2.1: Đặc trưng tổng quát của dữ liệu sau lấy mẫu",
    )
    add_figure(
        doc,
        "aggregate_sales_over_time.png",
        "Hình 2.1: Tổng doanh số theo thời gian trên tập dữ liệu mẫu",
        "Hinh_2.1_Tong_doanh_so_theo_thoi_gian.png",
    )
    add_figure(
        doc,
        "stockout_rate_over_time.png",
        "Hình 2.2: Tỷ lệ stockout theo thời gian",
        "Hinh_2.2_Ty_le_stockout_theo_thoi_gian.png",
    )
    add_paragraph(
        doc,
        "Tỷ lệ stockout trung bình đạt khoảng "
        f"{fmt_pct(quality.get('Stockout rate', ''))}, cho thấy đây không phải nhiễu nhỏ có thể bỏ qua. Nếu mô hình học trực tiếp trên observed sales, các giai đoạn hết hàng sẽ kéo target xuống thấp và làm mô hình hiểu nhầm rằng nhu cầu thực giảm. Đây là lý do nền tảng để bài toán được framing theo hướng latent demand recovery trước khi forecast.",
    )
    add_paragraph(
        doc,
        "Một điểm mù khác của observed sales là hiệu ứng thay thế sản phẩm. Khi sản phẩm A hết hàng, khách hàng có thể chuyển sang sản phẩm B cùng nhóm, làm sales của B tăng bất thường. Nếu chỉ nhìn từng chuỗi riêng lẻ, phần nhu cầu bị chuyển dịch này có thể bị bỏ qua. Do đó, các chương sau bổ sung tín hiệu peer/substitution theo cùng cửa hàng và cùng nhóm hàng để kiểm tra bối cảnh xung quanh stockout.",
    )

    doc.add_heading("2.2. Kiểm định Phân phối Thống kê của Doanh số", level=2)
    add_paragraph(
        doc,
        "Phân phối doanh số có đặc trưng lệch phải mạnh và tập trung nhiều tại 0. Trung vị doanh số là "
        f"{fmt_float(quality.get('Median sale_amount', ''), 4)}, trong khi trung bình là "
        f"{fmt_float(quality.get('Average sale_amount', ''), 4)}. Sự chênh lệch này phản ánh tính thưa của dữ liệu: phần lớn cặp cửa hàng-sản phẩm không phát sinh bán hàng trong nhiều giờ, nhưng vẫn tồn tại một số thời điểm có doanh số cao.",
    )
    add_figure(
        doc,
        "sale_amount_distribution.png",
        "Hình 2.3: Phân phối doanh số quan sát ở thang gốc",
        "Hinh_2.3_Phan_phoi_doanh_so_thang_goc.png",
    )
    add_figure(
        doc,
        "log_sale_amount_distribution.png",
        "Hình 2.4: Phân phối doanh số sau biến đổi log1p",
        "Hinh_2.4_Phan_phoi_doanh_so_log1p.png",
    )
    add_paragraph(
        doc,
        "Biến đổi log1p giúp quan sát phần đuôi phân phối rõ hơn nhưng không làm mất bản chất zero-inflated của dữ liệu. Đặc điểm này khiến MAPE kém phù hợp vì mẫu số có thể bằng 0 hoặc rất nhỏ. Vì vậy, các chỉ số như WAPE, MAE và RMSE phù hợp hơn cho phần đánh giá mô hình. Phân phối không chuẩn và nhiều chuỗi thưa cũng là lý do thực tế để ưu tiên global machine learning model thay vì fit riêng một mô hình classical time series cho từng chuỗi.",
    )
    add_paragraph(
        doc,
        "Như vậy, mục tiêu của phần này không phải chứng minh doanh số tuân theo một phân phối lý thuyết cố định, mà là chỉ ra rằng target có tính thưa, lệch và chịu kiểm duyệt bởi stockout. Điều này ảnh hưởng trực tiếp đến lựa chọn thước đo đánh giá, baseline và phương pháp recovery.",
    )

    doc.add_heading("2.3. Đánh giá Tính Mùa vụ (Seasonality) theo Chu kỳ Giờ và Tuần", level=2)
    add_paragraph(
        doc,
        "Dữ liệu bán lẻ thường chịu ảnh hưởng mạnh bởi nhịp sinh hoạt của khách hàng và lịch vận hành cửa hàng. Vì dữ liệu gốc có độ phân giải theo giờ, hai dạng mùa vụ quan trọng cần kiểm tra là mùa vụ trong ngày và mùa vụ trong tuần.",
    )
    add_figure(
        doc,
        "sales_by_hour_of_day.png",
        "Hình 2.5: Doanh số trung bình theo giờ trong ngày",
        "Hinh_2.5_Doanh_so_trung_binh_theo_gio.png",
    )
    add_figure(
        doc,
        "sales_by_day_of_week.png",
        "Hình 2.6: Doanh số trung bình theo ngày trong tuần",
        "Hinh_2.6_Doanh_so_trung_binh_theo_ngay_trong_tuan.png",
    )
    add_paragraph(
        doc,
        "Các biểu đồ mùa vụ cho thấy doanh số không phân bố đều theo thời gian mà thay đổi theo giờ và theo ngày trong tuần. Điều này biện minh cho các đặc trưng lag 24 giờ, 168 giờ ở cấp hourly, cũng như lag 7 ngày, 14 ngày và 28 ngày ở cấp daily. Kết quả này cũng giải thích vì sao seasonal naive là baseline mạnh: nếu hành vi mua hàng lặp lại theo tuần, dự báo dựa trên cùng kỳ tuần trước đã nắm được một phần đáng kể tín hiệu.",
    )
    add_paragraph(
        doc,
        "Riêng với stockout, seasonality theo giờ còn giúp hiểu điểm mù thứ hai của bài toán: tốc độ bán trong các giờ trước khi hết hàng có thể phản ánh demand đang tăng mạnh trước khi dữ liệu bị missing hoặc bị censor. Vì vậy, ở bước feature engineering, đồ án bổ sung các biến velocity/momentum như tỷ lệ rolling 3 giờ so với rolling 24 giờ và chênh lệch giữa các cửa sổ ngắn hạn.",
    )

    doc.add_heading("2.4. Lựa chọn và Đánh giá các Chuỗi Dữ liệu Đại diện", level=2)
    add_paragraph(
        doc,
        "Chuỗi tổng hợp có thể che khuất sự khác biệt giữa các nhóm cửa hàng-sản phẩm. Vì vậy, đồ án chọn một số chuỗi đại diện theo ba nhóm: high-volume, intermittent và stockout-heavy. Cách chọn này giúp đánh giá bài toán trên các tình huống có ý nghĩa kinh doanh khác nhau thay vì chỉ nhìn trung bình tổng thể.",
    )
    rep_path = TABLE_DIR / "representative_series.csv"
    if rep_path.exists():
        rep = pd.read_csv(rep_path).head(10)
        keep_cols = [col for col in rep.columns[:6]]
        add_table(
            doc,
            [str(col) for col in keep_cols],
            [[str(row[col]) for col in keep_cols] for _, row in rep.iterrows()],
            "Bảng 2.2: Một số chuỗi đại diện được chọn để phân tích trực quan",
        )
    add_figure(
        doc,
        "representative_series_high_volume.png",
        "Hình 2.7: Chuỗi đại diện nhóm doanh số cao",
        "Hinh_2.7_Chuoi_dai_dien_nhom_doanh_so_cao.png",
    )
    add_figure(
        doc,
        "representative_series_intermittent.png",
        "Hình 2.8: Chuỗi đại diện nhóm bán gián đoạn",
        "Hinh_2.8_Chuoi_dai_dien_nhom_ban_gian_doan.png",
    )
    add_figure(
        doc,
        "representative_series_stockout_heavy.png",
        "Hình 2.9: Chuỗi đại diện nhóm chịu stockout nặng",
        "Hinh_2.9_Chuoi_dai_dien_nhom_stockout_nang.png",
    )
    add_paragraph(
        doc,
        "Nhóm high-volume thường có tín hiệu rõ và dễ học mùa vụ hơn. Nhóm intermittent có nhiều giá trị 0, khiến dự báo từng thời điểm khó hơn và cần đánh giá ở cấp tổng hợp. Nhóm stockout-heavy là nhóm quan trọng để chứng minh đóng góp của đồ án: nếu không xử lý stockout, mô hình có nguy cơ học theo doanh số bị kiểm duyệt thay vì nhu cầu thực.",
    )
    add_paragraph(
        doc,
        "Sự khác biệt giữa các nhóm chuỗi là lý do chính để dùng global model. Mô hình toàn cục có thể học pattern chung giữa nhiều chuỗi, đồng thời dùng định danh cửa hàng, sản phẩm và nhóm hàng để phân biệt hành vi riêng, thay vì huấn luyện thủ công một mô hình riêng cho từng chuỗi thưa.",
    )

    doc.add_heading("2.5. Phân tích Đặc tính Chuỗi Thời gian: Kiểm định Tính Dừng (Stationary Testing)", level=2)
    add_paragraph(
        doc,
        "Tính dừng được kiểm tra bằng ADF và KPSS như một bước chẩn đoán, không phải để ép toàn bộ bài toán về ARIMA/SARIMA. ADF có giả thuyết gốc là chuỗi không dừng, trong khi KPSS có giả thuyết gốc là chuỗi dừng. Khi hai kiểm định cho kết quả khác nhau, kết luận nên được xem là bằng chứng hỗn hợp.",
    )
    add_stationarity_table(doc)
    add_paragraph(
        doc,
        "Kết quả cho thấy chuỗi tổng hợp ở dạng gốc có dấu hiệu không ổn định, nhưng sau sai phân hoặc sai phân mùa vụ thì trở nên dừng rõ hơn. Điều này xác nhận rằng dữ liệu có cấu trúc phụ thuộc theo thời gian và seasonality. Tuy nhiên, vì dữ liệu gồm nhiều chuỗi thưa, nhiều stockout và nhiều covariates, việc fit một mô hình classical riêng cho từng chuỗi sẽ khó mở rộng và khó ổn định.",
    )
    add_paragraph(
        doc,
        "Do đó, kiểm định tính dừng được dùng để thiết kế đặc trưng như lag, rolling mean, rolling sum và seasonal lag, còn mô hình chính vẫn là global ML model. Cách tiếp cận này giữ được tinh thần phân tích chuỗi thời gian nhưng phù hợp hơn với quy mô và bản chất dữ liệu bán lẻ nhiều chuỗi.",
    )

    doc.add_heading("2.6. Tự tương quan (ACF/PACF) và Phân tích Phổ (Spectrum Analysis)", level=2)
    add_paragraph(
        doc,
        "ACF/PACF được dùng để kiểm tra trực tiếp cấu trúc phụ thuộc theo lag. Nếu ACF duy trì ở các lag mùa vụ, mô hình nên được cung cấp các biến trễ tương ứng. PACF giúp quan sát phần phụ thuộc trực tiếp sau khi loại bớt ảnh hưởng của các lag trung gian. Trong bài này, ACF/PACF không nhằm chọn bậc ARIMA cuối cùng, mà đóng vai trò xác nhận các lag feature cho mô hình học máy.",
    )
    add_figure(
        doc,
        "acf_aggregate_sales.png",
        "Hình 2.10: ACF của chuỗi doanh số tổng hợp",
        "Hinh_2.10_ACF_chuoi_doanh_so_tong_hop.png",
    )
    add_figure(
        doc,
        "pacf_aggregate_sales.png",
        "Hình 2.11: PACF của chuỗi doanh số tổng hợp",
        "Hinh_2.11_PACF_chuoi_doanh_so_tong_hop.png",
    )
    add_paragraph(
        doc,
        "Phân tích phổ chuyển chuỗi thời gian sang miền tần số để xác định các chu kỳ nổi bật. Kết quả ở cấp giờ cho thấy các chu kỳ ngắn trong ngày và dấu hiệu chu kỳ gần tuần. Ở cấp ngày sau recovery, chu kỳ xấp xỉ 7 ngày nổi bật, củng cố vai trò của seasonal naive 7-day và các đặc trưng lag theo tuần.",
    )
    add_figure(
        doc,
        "spectrum_hourly_aggregate_sales.png",
        "Hình 2.12: Phổ tần số của chuỗi doanh số tổng hợp theo giờ",
        "Hinh_2.12_Pho_tan_so_doanh_so_theo_gio.png",
    )
    add_spectrum_table(
        doc,
        "spectrum_hourly_top_peaks.csv",
        "Bảng 2.4: Các chu kỳ nổi bật trong phân tích phổ ở cấp giờ",
        "giờ",
    )
    add_figure(
        doc,
        "spectrum_daily_recovered_demand.png",
        "Hình 2.13: Phổ tần số của chuỗi nhu cầu phục hồi ở cấp ngày",
        "Hinh_2.13_Pho_tan_so_nhu_cau_phuc_hoi_theo_ngay.png",
    )
    add_spectrum_table(
        doc,
        "spectrum_daily_recovered_top_peaks.csv",
        "Bảng 2.5: Các chu kỳ nổi bật trong phân tích phổ ở cấp ngày",
        "ngày",
    )
    add_paragraph(
        doc,
        "Như vậy, ACF/PACF và spectrum xác nhận chéo cho kết quả seasonality ở mục 2.3. Điều này giúp việc dùng seasonal baseline, lag theo tuần và hybrid seasonal-ML ở các chương sau có cơ sở từ cả trực quan dữ liệu lẫn phân tích tần số.",
    )

    if sub_velocity_rows:
        doc.add_heading("2.7. Tín hiệu Velocity và Peer/Substitution quanh Stockout", level=2)
        add_paragraph(
            doc,
            "Sau khi xác định stockout là vấn đề trung tâm, đồ án kiểm tra thêm hai tín hiệu có ý nghĩa nghiệp vụ. Tín hiệu velocity/momentum đo xem sản phẩm có bán nhanh hơn nền lịch sử ngay trước hoặc trong giai đoạn stockout hay không. Tín hiệu peer/substitution đo xem các sản phẩm cùng cửa hàng và cùng nhóm hàng có bán tăng lên hoặc cùng gặp stockout hay không.",
        )
        add_table(
            doc,
            [
                "Segment",
                "Rows",
                "Velocity 3h/24h",
                "Momentum 3h-6h",
                "Peer sales",
                "Peer velocity 3h/24h",
                "Peer stockout rate",
            ],
            sub_velocity_rows,
            "Bảng 2.6: So sánh tín hiệu velocity và peer/substitution giữa stockout và non-stockout",
        )
        add_figure(
            doc,
            "stockout_substitution_velocity_diagnostics.png",
            "Hình 2.14: Tín hiệu velocity và peer/substitution quanh stockout",
            "Hinh_2.14_Tin_hieu_velocity_peer_stockout.png",
        )
        add_paragraph(
            doc,
            "Bảng và hình trên không nhằm khẳng định mọi stockout đều đi kèm demand surge. Ngược lại, chúng giúp kiểm tra có hệ thống hai điểm mù của observed sales: nhu cầu có thể tăng trước khi mất quan sát và nhu cầu có thể chuyển sang sản phẩm thay thế. Vì vậy, các feature velocity và peer/substitution được đưa vào giai đoạn recovery để mô hình có thêm bối cảnh khi ước lượng latent demand.",
        )

    doc.add_heading("Kết luận Chương 2", level=2)
    add_paragraph(
        doc,
        "Chương 2 cho thấy bài toán dự báo trên FreshRetailNet-50K không nên được xem như bài toán forecast doanh số quan sát thông thường. Observed sales vừa thưa, vừa lệch, vừa có thể bị kiểm duyệt bởi stockout. Đồng thời, dữ liệu có seasonality rõ theo giờ và theo tuần, nên các baseline mùa vụ và lag features là bắt buộc để đánh giá mô hình một cách công bằng.",
    )
    add_paragraph(
        doc,
        "Các phân tích về chuỗi đại diện, tính dừng, ACF/PACF, spectrum, velocity và peer/substitution dẫn đến framing cuối cùng của đồ án: phục hồi latent demand ở cấp giờ bằng cơ chế tránh leakage, tổng hợp lên daily demand, sau đó forecast next-7-day demand bằng seasonal baseline, LightGBM và hybrid seasonal-ML. Như vậy, Chương 2 là cơ sở thực nghiệm trực tiếp cho phương pháp ở Chương 3 và diễn giải kết quả ở Chương 4.",
    )

    saved_path = save_doc(doc, OUT_PATH)
    print(saved_path)


if __name__ == "__main__":
    build_doc()
